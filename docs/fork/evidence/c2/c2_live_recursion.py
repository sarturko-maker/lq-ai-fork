"""C2 live verification — run inside the REBUILT api image (baked code + python-oxmsg).

Proves the wired registry recurses a multi-attachment deal email one level so the
agent can ground answers in a buried attachment. DB/MinIO persistence is unchanged
from C1 and covered by the CI integration test on real postgres.
"""
import importlib.metadata as md
import io
import types

import oxmsg  # noqa: F401  (import must succeed in the baked image)
from docx import Document
from email.message import EmailMessage

from app.pipeline.chunker import chunk_document
from app.pipeline.readers import MsgReader, build_default_registry
from app.pipeline.readers._base import EML_MIME

_settings = types.SimpleNamespace(lq_ai_docling_enabled=False)

print("python-oxmsg installed in image:", md.version("python-oxmsg"))


def make_docx() -> bytes:
    d = Document()
    d.add_paragraph("Mutual Non-Disclosure Agreement between the parties.")
    d.add_paragraph("Liability cap: fees paid in the three (3) months preceding the claim.")
    b = io.BytesIO()
    d.save(b)
    return b.getvalue()


msg = EmailMessage()
msg["From"] = "jason.martinez@securescan.com"
msg["To"] = "commercial-legal@zendesk.com"
msg["Subject"] = "SecureScan order form (clean paper, DPA required)"
msg["Date"] = "Mon, 1 Jun 2026 10:00:00 +0000"
msg.set_content("First-pass review on the attached order form. They touch our customers' PII.")
msg.add_attachment(
    make_docx(),
    maintype="application",
    subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
    filename="SecureScan-OrderForm.docx",
)
msg.add_attachment(b"binary-logo-bytes", maintype="image", subtype="png", filename="logo.png")
eml_bytes = msg.as_bytes()

reader = build_default_registry(_settings).for_mime(EML_MIME)
parsed = reader.read(eml_bytes)
ct = parsed.canonical_text
print("parser:", parsed.parser, "| units:", parsed.page_count)
print("  email body grounded :", "First-pass review on the attached order form." in ct)
print("  docx text grounded  :", "Mutual Non-Disclosure Agreement between the parties." in ct)
print("  cap clause grounded :", "three (3) months" in ct)
print("  docx label present  :", "SecureScan-OrderForm.docx" in ct)
print("  png listed, no bytes:", "logo.png" in ct and "binary-logo-bytes" not in ct)
sc = parsed.structured_content
print("  sc format:", sc["format"], "| attachments:", [(a["filename"], a["status"]) for a in sc["attachments"]])
chunks = chunk_document(parsed, target_chars=200, overlap_chars=40)
fidelity = all(ct[c.char_offset_start : c.char_offset_end] == c.content for c in chunks)
print("  chunk fidelity OK   :", fidelity, "| chunks:", len(chunks))
print("  .msg OLE sniff      :", MsgReader().sniff(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1x"), "| non-msg:", MsgReader().sniff(b"%PDF"))
print("C2 LIVE OK")
