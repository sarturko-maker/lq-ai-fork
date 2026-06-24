"""C4 golden-redline corpus — Layer 1 + Layer 2 (ADR-F031).

The maintainer's requirement: judge the PRODUCED ``.docx``, not the word count,
and separate *mechanism* quality (does the system render a good redline from good
edits) from *model* quality. So this is **model-free**: each scenario carries a
hand-authored *known-good surgical redline* (the §5.1 edits a senior lawyer
makes) and we assert, against the actually-produced document:

* the edits pass the surgical gate (Layer 1),
* the rendering is surgical — unchanged head/tail/interior stay BARE, only the
  changed spans are tracked (Layer 2, read from ``word/document.xml``),
* accepting all changes yields the balanced clause: the protective language is
  present and the one-sided language is gone.

These need Adeu + python-docx (both pinned deps); they run in CI. The same GOLDEN
corpus is the gate's calibration corpus and the live-evidence input.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field

import pytest

from app.agents.redline_render import bare_text, docx_text, reconstruct_redline_text
from app.agents.redline_service import ProposedEdit, RedlineService
from app.schemas.commercial import RedlineEditInput, evaluate_gate


@dataclass(frozen=True)
class GoldenScenario:
    name: str
    paragraphs: list[str]
    edits: list[dict[str, str]]
    expect_bare: list[str] = field(default_factory=list)  # stays UNCHANGED (surgical)
    expect_in_accepted: list[str] = field(default_factory=list)  # protection added
    forbidden_in_accepted: list[str] = field(default_factory=list)  # one-sided text gone


_R_PERIOD = (
    "The house liability floor is a twelve-month fee measure, so extend the lookback "
    "period to align the cap with the customer's standard position."
)
_R_CARVE = (
    "Carve the high-risk heads of loss out of the cap rather than merely raising the "
    "figure, so a confidentiality, data-protection or IP breach is not capped."
)
_R_MUTUAL = (
    "Make the indemnity mutual so each party bears responsibility for its own "
    "wrongdoing; a one-way indemnity in the counterparty's favour is off-market here."
)
_R_IP = (
    "Preserve each party's pre-existing and independently developed IP and grant only "
    "a licence to background materials, rather than assigning the vendor's own IP away."
)
_R_IP_RECIP = (
    "Add a reciprocal indemnity so the vendor stands behind third-party intellectual "
    "property infringement claims arising from its own materials, balancing the "
    "one-sided original that protected only the vendor."
)


GOLDEN: list[GoldenScenario] = [
    GoldenScenario(
        name="vendor-favoured-limitation-of-liability",
        paragraphs=[
            "The Vendor's aggregate liability arising out of or in connection with this "
            "Agreement shall not exceed the total fees paid by the Customer in the three "
            "(3) months preceding the claim."
        ],
        edits=[
            {
                "target_text": "three (3)",
                "new_text": "twelve (12)",
                "rationale": _R_PERIOD,
            },
            {
                "target_text": "preceding the claim.",
                "new_text": (
                    "preceding the claim, save that liability for breach of "
                    "confidentiality, data protection obligations or infringement of "
                    "intellectual property rights shall be unlimited."
                ),
                "rationale": _R_CARVE,
            },
        ],
        expect_bare=[
            "shall not exceed the total fees paid by the Customer in the",
            "months preceding the",
        ],
        expect_in_accepted=["twelve (12) months", "shall be unlimited"],
        forbidden_in_accepted=["three (3) months"],
    ),
    GoldenScenario(
        name="one-sided-indemnity-mutualised",
        paragraphs=[
            "The Customer shall indemnify, defend and hold harmless the Vendor and its "
            "affiliates against any and all claims, losses, damages and liabilities "
            "arising from or in connection with the performance of this Agreement."
        ],
        edits=[
            {
                "target_text": "The Customer shall indemnify, defend and hold harmless the Vendor",
                "new_text": "Each party shall indemnify, defend and hold harmless the other",
                "rationale": _R_MUTUAL,
            },
        ],
        expect_bare=["against any and all claims, losses, damages and liabilities"],
        expect_in_accepted=["Each party shall indemnify"],
        forbidden_in_accepted=["The Customer shall indemnify"],
    ),
    GoldenScenario(
        name="one-sided-ip-assignment-carved",
        paragraphs=[
            "All intellectual property rights in any deliverables, materials and work "
            "product created under this Agreement shall vest exclusively in the Customer "
            "upon creation."
        ],
        edits=[
            {
                "target_text": "vest exclusively in the Customer upon creation.",
                "new_text": (
                    "vest exclusively in the Customer upon creation, save that each "
                    "party retains all intellectual property rights it owned prior to "
                    "or independently of this Agreement, and the Vendor grants the "
                    "Customer a non-exclusive licence to its background materials to "
                    "the extent necessary to use the deliverables."
                ),
                "rationale": _R_IP,
            },
        ],
        expect_bare=["vest exclusively in the Customer upon"],
        expect_in_accepted=["non-exclusive licence", "background materials"],
        forbidden_in_accepted=[],
    ),
    # C8 (ADR-F041): the §8 mutualisation done SURGICALLY — several narrow edits
    # (two party swaps + one inserted reciprocal indemnity), each its own
    # ModifyText, so the verb-phrase boilerplate "shall indemnify, defend and hold
    # harmless" stays BARE. The C8 craft the surgical-redline skill teaches; the
    # anti-pattern is one clause-sized strike-and-retype.
    GoldenScenario(
        name="indemnity-mutualised-surgically",
        paragraphs=[
            "The Customer shall indemnify, defend and hold harmless the Vendor and its "
            "affiliates against any and all claims, losses, damages and liabilities "
            "arising from or in connection with the performance of this Agreement."
        ],
        edits=[
            {"target_text": "The Customer", "new_text": "Each party", "rationale": _R_MUTUAL},
            {
                "target_text": "the Vendor and its affiliates",
                "new_text": "the other party and its affiliates",
                "rationale": _R_MUTUAL,
            },
            {
                # Insertion folded into the boundary (replace the closing ".") — a
                # pure zero-width append after an unchanged anchor crashes the editor.
                "target_text": "the performance of this Agreement.",
                "new_text": (
                    "the performance of this Agreement, and each party shall indemnify "
                    "the other against third-party intellectual property infringement claims."
                ),
                "rationale": _R_IP_RECIP,
            },
        ],
        expect_bare=[
            "shall indemnify, defend and hold harmless",  # boilerplate never touched
            "against any and all claims, losses, damages and liabilities",
        ],
        expect_in_accepted=[
            "Each party shall indemnify",
            "third-party intellectual property infringement",
        ],
        forbidden_in_accepted=["The Customer shall indemnify"],
    ),
]


def build_docx(paragraphs: list[str]) -> bytes:
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def apply_golden(service: RedlineService, scenario: GoldenScenario) -> bytes:
    """Run a scenario's known-good edits through the real apply path."""
    logical = [
        ProposedEdit(e["target_text"], e["new_text"], e.get("rationale") or None)
        for e in scenario.edits
    ]
    docx = build_docx(scenario.paragraphs)
    preview = service.dry_run(docx, logical)
    assert preview.edits_skipped == 0 and preview.edits_applied > 0, (
        f"{scenario.name}: applied={preview.edits_applied} skipped={preview.edits_skipped}"
    )
    return service.apply(docx, logical).docx_bytes


@pytest.mark.parametrize("scenario", GOLDEN, ids=lambda s: s.name)
def test_golden_edits_pass_the_surgical_gate(scenario: GoldenScenario) -> None:
    """Layer 1: the hand-authored known-good edits clear the gate."""
    document_text = "\n".join(scenario.paragraphs)
    edits = [RedlineEditInput(**e) for e in scenario.edits]  # type: ignore[arg-type]
    report = evaluate_gate(document_text, edits)
    assert report.ok, f"{scenario.name}: {report.rejection_text()}"


@pytest.mark.parametrize("scenario", GOLDEN, ids=lambda s: s.name)
def test_golden_redline_renders_surgically(scenario: GoldenScenario) -> None:
    """Layer 2: read the produced .docx — unchanged words stay bare, only the
    changed spans are tracked, and accept-to-clean is the balanced clause."""
    service = RedlineService()
    redlined = apply_golden(service, scenario)

    redline = reconstruct_redline_text(redlined)
    assert "[+" in redline or "[-" in redline, "no tracked changes rendered"

    bare = bare_text(redline)
    for phrase in scenario.expect_bare:
        assert phrase in bare, (
            f"{scenario.name}: '{phrase}' should stay BARE (unchanged) but was swept "
            f"into a tracked change — over-marking. Redline:\n{redline}"
        )

    clean = docx_text(service.accept_all(redlined))
    for phrase in scenario.expect_in_accepted:
        assert phrase in clean, f"{scenario.name}: accepted text missing '{phrase}'"
    for phrase in scenario.forbidden_in_accepted:
        assert phrase not in clean, f"{scenario.name}: accepted text still has '{phrase}'"
