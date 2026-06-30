"""Matter document tools — F0-S4: FTS search + full read over REAL rows.

Real Postgres, real ``content_tsv`` FTS (the migrated per-run test DB),
real audit rows: ``build_matter_tools`` closures are exercised exactly
as the runner dispatches them. The load-bearing assertions:

* scoping — only the bound matter's documents, only the run owner's
  files (a directly-seeded foreign-owner ``project_files`` row — a
  state the API can't produce — must still be excluded: defense in
  depth, the chats-path posture);
* honesty — inventory on empty query, explicit no-match message,
  ingestion-pending notice, bounded read with a truncation notice;
* the guard chokepoint — every dispatch leaves exactly one
  ``agent_run.tool_call`` audit row whose details carry counts/types/
  IDs, never document content (CLAUDE.md audit contract).

These tests COMMIT (the tools open their own sessions from the
factory), so they seed via a commit-capable factory and tear down
explicitly — FK order: audit rows are SET NULL, projects/files are
RESTRICT on owner, so children go first.
"""

from __future__ import annotations

import inspect
import io
import uuid
from collections.abc import AsyncIterator, Awaitable, Callable
from contextlib import asynccontextmanager
from dataclasses import dataclass
from types import SimpleNamespace

import pytest
import pytest_asyncio
from sqlalchemy import delete, select
from sqlalchemy.ext.asyncio import AsyncEngine, AsyncSession, async_sessionmaker

from app.agents import tools as agent_tools
from app.agents.tools import MATTER_TOOL_NAMES, MatterBinding, build_matter_tools
from app.knowledge import rerank_provider as rp
from app.models.agent_run import AgentRun, AgentThread
from app.models.audit import AuditLog
from app.models.document import Document, DocumentChunk
from app.models.file import File
from app.models.project import Project, ProjectFile
from app.models.user import User
from app.pipeline.readers._base import OOXML_DOCX_MIME
from app.security import hash_password

pytestmark = pytest.mark.integration

_MSA_TEXT = (
    "Clause 7.2 (Limitation of Liability): each party's aggregate liability "
    "is capped at the fees paid in the twelve (12) months preceding the claim."
)
_NOTES_TEXT = "Meeting notes: timeline review, renewal checkpoints, owner actions."
_UPLOADED_TEXT = "Escrow arrangements for the source code deposit with Iron Mountain."
_SECRET_TEXT = "Confidential merger arrangement discussions with Northwind."
_FOREIGN_TEXT = "Foreign-owner dossier: zugzwang acquisition memorandum."


@dataclass
class MatterEnv:
    factory: async_sessionmaker[AsyncSession]
    user_id: uuid.UUID
    run_id: uuid.UUID
    project_id: uuid.UUID
    search: Callable[..., Awaitable[str]]
    read: Callable[..., Awaitable[str]]
    get_meta: Callable[..., Awaitable[str]]
    estimate: Callable[..., Awaitable[str]]


@pytest_asyncio.fixture
async def commit_factory(test_engine: AsyncEngine) -> async_sessionmaker[AsyncSession]:
    return async_sessionmaker(bind=test_engine, expire_on_commit=False, class_=AsyncSession)


async def _seed_file(
    db: AsyncSession,
    *,
    owner_id: uuid.UUID,
    filename: str,
    body: str | None,
    page_start: int | None = 1,
    page_end: int | None = 1,
    column_project_id: uuid.UUID | None = None,
) -> File:
    """File (+ Document + one chunk when ``body`` is given; None = pending).

    ``column_project_id`` sets the upload-time ``files.project_id``
    column (the ``POST /files`` affordance) — membership without a
    ``project_files`` join row.
    """
    f = File(
        owner_id=owner_id,
        project_id=column_project_id,
        filename=filename,
        mime_type="application/pdf",
        size_bytes=1024,
        hash_sha256="a" * 64,
        storage_path=f"agent-tools-fixture/{uuid.uuid4()}",
        ingestion_status="processing" if body is None else "ready",
    )
    db.add(f)
    await db.flush()
    if body is not None:
        doc = Document(
            file_id=f.id,
            parser="pymupdf-only",
            page_count=page_end or 1,
            character_count=len(body),
            normalized_content=body,
        )
        db.add(doc)
        await db.flush()
        db.add(
            DocumentChunk(
                document_id=doc.id,
                chunk_index=0,
                content=body,
                page_start=page_start,
                page_end=page_end,
                char_offset_start=0,
                char_offset_end=len(body),
            )
        )
    return f


@pytest_asyncio.fixture
async def matter_env(
    commit_factory: async_sessionmaker[AsyncSession],
) -> AsyncIterator[MatterEnv]:
    """One matter with two ingested docs + one pending doc, plus decoys:

    * a second matter (same owner) whose document must never surface;
    * a foreign-owned file maliciously joined into OUR matter — the
      owner re-assertion must exclude it.
    """
    user_ids: list[uuid.UUID] = []
    project_ids: list[uuid.UUID] = []
    file_ids: list[uuid.UUID] = []

    async with commit_factory() as db:
        user = User(
            email=f"agent-tools-{uuid.uuid4().hex[:8]}@example.com",
            display_name="Agent Tools User",
            hashed_password=hash_password("correct-horse-battery-staple"),
            is_admin=False,
            mfa_enabled=False,
            must_change_password=False,
        )
        stranger = User(
            email=f"agent-tools-stranger-{uuid.uuid4().hex[:8]}@example.com",
            display_name="Agent Tools Stranger",
            hashed_password=hash_password("correct-horse-battery-staple"),
            is_admin=False,
            mfa_enabled=False,
            must_change_password=False,
        )
        db.add_all([user, stranger])
        await db.flush()
        user_ids += [user.id, stranger.id]

        project = Project(owner_id=user.id, name="Acme MSA", slug=f"acme-{uuid.uuid4().hex[:6]}")
        other_project = Project(
            owner_id=user.id, name="Northwind", slug=f"north-{uuid.uuid4().hex[:6]}"
        )
        db.add_all([project, other_project])
        await db.flush()
        project_ids += [project.id, other_project.id]

        msa = await _seed_file(
            db,
            owner_id=user.id,
            filename="msa.pdf",
            body=_MSA_TEXT,
            page_start=7,
            page_end=8,
        )
        notes = await _seed_file(db, owner_id=user.id, filename="notes.pdf", body=_NOTES_TEXT)
        pending = await _seed_file(db, owner_id=user.id, filename="pending.pdf", body=None)
        # Upload-time membership ONLY (files.project_id column, no join
        # row) — exactly what POST /files with project_id produces.
        uploaded = await _seed_file(
            db,
            owner_id=user.id,
            filename="uploaded.pdf",
            body=_UPLOADED_TEXT,
            column_project_id=project.id,
        )
        secret = await _seed_file(db, owner_id=user.id, filename="secret.pdf", body=_SECRET_TEXT)
        # Foreign-owned, with the column ALSO spoofed at our matter — the
        # owner re-assertion must exclude it through both membership paths.
        foreign = await _seed_file(
            db,
            owner_id=stranger.id,
            filename="foreign.pdf",
            body=_FOREIGN_TEXT,
            column_project_id=project.id,
        )
        file_ids += [msa.id, notes.id, pending.id, uploaded.id, secret.id, foreign.id]

        db.add_all(
            [
                ProjectFile(project_id=project.id, file_id=msa.id),
                ProjectFile(project_id=project.id, file_id=notes.id),
                ProjectFile(project_id=project.id, file_id=pending.id),
                ProjectFile(project_id=other_project.id, file_id=secret.id),
                # The attack row: a foreign-owned file joined into OUR
                # matter. The API can't create this (attach requires
                # owning both); the tools must still exclude it.
                ProjectFile(project_id=project.id, file_id=foreign.id),
            ]
        )

        thread = AgentThread(user_id=user.id, project_id=project.id, title="tools tests")
        db.add(thread)
        await db.flush()
        run = AgentRun(
            user_id=user.id,
            thread_id=thread.id,
            project_id=project.id,
            status="running",
            prompt="What is the liability cap?",
            model_alias="smart",
            max_steps=20,
        )
        db.add(run)
        await db.commit()

        binding = MatterBinding(
            project_id=project.id,
            user_id=user.id,
            name="Acme MSA",
            privileged=False,
            minimum_inference_tier=None,
        )
        search, read, get_meta, estimate = build_matter_tools(
            commit_factory, run_id=run.id, binding=binding
        )
        env = MatterEnv(
            factory=commit_factory,
            user_id=user.id,
            run_id=run.id,
            project_id=project.id,
            search=search,
            read=read,
            get_meta=get_meta,
            estimate=estimate,
        )

    yield env

    async with commit_factory() as db:
        await db.execute(delete(AuditLog).where(AuditLog.user_id.in_(user_ids)))
        await db.execute(delete(AgentRun).where(AgentRun.user_id.in_(user_ids)))
        await db.execute(delete(AgentThread).where(AgentThread.user_id.in_(user_ids)))
        await db.execute(delete(Project).where(Project.id.in_(project_ids)))
        await db.execute(delete(File).where(File.id.in_(file_ids)))
        await db.execute(delete(User).where(User.id.in_(user_ids)))
        await db.commit()


async def _audit_rows(env: MatterEnv) -> list[AuditLog]:
    async with env.factory() as db:
        rows = (
            await db.execute(
                select(AuditLog)
                .where(
                    AuditLog.resource_type == "agent_run",
                    AuditLog.resource_id == str(env.run_id),
                )
                .order_by(AuditLog.timestamp.asc(), AuditLog.id.asc())
            )
        ).scalars()
        return list(rows)


# ---------------------------------------------------------------------------
# search_documents
# ---------------------------------------------------------------------------


async def test_search_returns_passages_with_filename_and_pages(
    matter_env: MatterEnv,
) -> None:
    result = await matter_env.search("liability cap")
    assert "msa.pdf" in result
    assert "pages 7-8" in result
    assert "aggregate liability" in result
    # The other matter's and the foreign-owned content never surface.
    assert "merger" not in result and "zugzwang" not in result


async def test_search_excludes_other_matters_of_the_same_owner(
    matter_env: MatterEnv,
) -> None:
    result = await matter_env.search("confidential merger Northwind")
    assert "No passages matched" in result
    assert "secret.pdf" not in result


async def test_search_excludes_foreign_owned_files_joined_into_the_matter(
    matter_env: MatterEnv,
) -> None:
    result = await matter_env.search("zugzwang acquisition")
    assert "No passages matched" in result
    assert "foreign.pdf" not in result


async def test_search_empty_query_lists_the_inventory(matter_env: MatterEnv) -> None:
    result = await matter_env.search("")
    assert "msa.pdf" in result and "notes.pdf" in result
    assert "pending.pdf" in result and "not ingested yet" in result
    assert "uploaded.pdf" in result  # column-only membership counts
    assert "foreign.pdf" not in result
    # F2 Slice E: each ingested doc carries a read-cost estimate (~k tokens to read);
    # the not-ingested file has no Document, so no estimate (just the status notice).
    assert "tokens to read" in result
    pending_line = next(line for line in result.splitlines() if "pending.pdf" in line)
    assert "tokens to read" not in pending_line


async def test_search_no_hits_is_honest_and_orients_the_model(
    matter_env: MatterEnv,
) -> None:
    result = await matter_env.search("hovercraft eels")
    assert 'No passages matched "hovercraft eels"' in result
    assert "msa.pdf" in result  # inventory so the model can pivot to read_document


async def test_search_finds_upload_time_membership(matter_env: MatterEnv) -> None:
    """POST /files with project_id sets only files.project_id (no join
    row) — those documents are still the matter's (verified live)."""
    result = await matter_env.search("escrow source code deposit")
    assert "uploaded.pdf" in result
    assert "Iron Mountain" in result


async def test_search_rerank_enabled_routes_through_the_reranker(
    matter_env: MatterEnv, monkeypatch: pytest.MonkeyPatch
) -> None:
    """With ``rerank_enabled`` the tool routes through the cross-encoder path; the
    autouse identity reranker preserves order, so the expected passage still surfaces
    (Slice D wiring, end-to-end through the guarded tool)."""
    monkeypatch.setattr(
        agent_tools,
        "get_settings",
        lambda: SimpleNamespace(rerank_enabled=True, rerank_candidates=30),
    )
    result = await matter_env.search("liability cap")
    assert "msa.pdf" in result and "aggregate liability" in result


async def test_search_rerank_error_degrades_to_hybrid(
    matter_env: MatterEnv, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A reranker failure never breaks the tool — it falls back to the hybrid order
    (the never-hard-fail posture, mirroring the embedder fallback)."""
    monkeypatch.setattr(
        agent_tools,
        "get_settings",
        lambda: SimpleNamespace(rerank_enabled=True, rerank_candidates=30),
    )
    from tests.agents.embedding_fakes import KeywordRerankProvider

    rp.set_rerank_provider(KeywordRerankProvider(fail=True))
    result = await matter_env.search("liability cap")
    assert "msa.pdf" in result and "aggregate liability" in result


# ---------------------------------------------------------------------------
# read_document
# ---------------------------------------------------------------------------


async def test_read_document_works_for_upload_time_membership(
    matter_env: MatterEnv,
) -> None:
    result = await matter_env.read("uploaded.pdf")
    assert "full text" in result
    assert _UPLOADED_TEXT in result


async def test_read_document_returns_full_text_case_insensitive(
    matter_env: MatterEnv,
) -> None:
    result = await matter_env.read("MSA.PDF")
    assert "full text" in result
    assert _MSA_TEXT in result


async def test_read_document_unknown_name_lists_inventory(
    matter_env: MatterEnv,
) -> None:
    result = await matter_env.read("ghost.pdf")
    assert 'No document named "ghost.pdf"' in result
    assert "msa.pdf" in result


async def test_read_document_foreign_owned_file_is_invisible(
    matter_env: MatterEnv,
) -> None:
    result = await matter_env.read("foreign.pdf")
    assert 'No document named "foreign.pdf"' in result
    assert "zugzwang" not in result


async def test_read_document_pending_ingestion_is_honest(matter_env: MatterEnv) -> None:
    result = await matter_env.read("pending.pdf")
    assert "no extractable text yet" in result


async def test_read_document_truncates_long_documents(matter_env: MatterEnv) -> None:
    long_body = "lorem ipsum dolor sit amet " * 2000  # ~54k chars
    async with matter_env.factory() as db:
        f = await _seed_file(db, owner_id=matter_env.user_id, filename="tome.pdf", body=long_body)
        db.add(ProjectFile(project_id=matter_env.project_id, file_id=f.id))
        await db.commit()
        tome_id = f.id

    try:
        result = await matter_env.read("tome.pdf")
        assert "more truncated" in result
        assert "search_documents" in result  # steers the model back to search
        assert len(result) < len(long_body)
    finally:
        async with matter_env.factory() as db:
            await db.execute(delete(File).where(File.id == tome_id))
            await db.commit()


async def test_read_document_duplicates_prefer_newest_readable(
    matter_env: MatterEnv,
) -> None:
    """Duplicate filenames resolve to the most recently ADDED readable
    copy — attach time when a join row exists, upload time for
    column-only members (which have no attached_at; F0-S4 review) —
    and an unreadable (pending) newest copy never shadows a readable one."""
    from datetime import UTC, datetime, timedelta

    now = datetime.now(UTC)
    async with matter_env.factory() as db:
        old_attached = await _seed_file(
            db,
            owner_id=matter_env.user_id,
            filename="dup.pdf",
            body="OLD attached copy.",
        )
        new_uploaded = await _seed_file(
            db,
            owner_id=matter_env.user_id,
            filename="dup.pdf",
            body="NEW uploaded copy.",
            column_project_id=matter_env.project_id,
        )
        pending_newest = await _seed_file(
            db,
            owner_id=matter_env.user_id,
            filename="dup.pdf",
            body=None,  # not ingested — unreadable
            column_project_id=matter_env.project_id,
        )
        old_attached.created_at = now - timedelta(days=3)
        new_uploaded.created_at = now - timedelta(days=1)
        pending_newest.created_at = now
        db.add(
            ProjectFile(
                project_id=matter_env.project_id,
                file_id=old_attached.id,
                attached_at=now - timedelta(days=2),
            )
        )
        await db.commit()
        dup_ids = [old_attached.id, new_uploaded.id, pending_newest.id]

    try:
        result = await matter_env.read("dup.pdf")
        assert "3 files in this matter share this name" in result
        assert "NEW uploaded copy." in result  # readable beats pending; newest wins
        assert "OLD attached copy." not in result
    finally:
        async with matter_env.factory() as db:
            await db.execute(delete(File).where(File.id.in_(dup_ids)))
            await db.commit()


# ---------------------------------------------------------------------------
# get_document_metadata (ADR-F048 Slice 2) — authorship attribution
# ---------------------------------------------------------------------------


async def test_get_document_metadata_email_returns_headers(matter_env: MatterEnv) -> None:
    """An email's stored structured_content headers (From/To/Date/Subject) are surfaced."""
    async with matter_env.factory() as db:
        f = File(
            owner_id=matter_env.user_id,
            project_id=matter_env.project_id,  # column membership
            filename="deal.eml",
            mime_type="message/rfc822",
            size_bytes=10,
            hash_sha256="b" * 64,
            storage_path=f"meta-fixture/{uuid.uuid4()}",
            ingestion_status="ready",
        )
        db.add(f)
        await db.flush()
        db.add(
            Document(
                file_id=f.id,
                parser="eml",
                page_count=1,
                character_count=5,
                normalized_content="From: ...",
                structured_content={
                    "format": "email",
                    "message_count": 1,
                    "messages": [
                        {
                            "ordinal": 1,
                            "type": "message",
                            "from": "Mark Counsel <mcounsel@beta.example>",
                            "to": "us@firm.example",
                            "date": "2026-06-26",
                            "subject": "Revised NDA",
                        }
                    ],
                },
            )
        )
        await db.commit()
        eid = f.id
    try:
        out = await matter_env.get_meta("deal.eml")
        assert "EMAIL METADATA" in out
        assert "From: Mark Counsel <mcounsel@beta.example>" in out
        assert "Subject: Revised NDA" in out
        assert "forgeable" in out.lower()  # untrusted-input framing
    finally:
        async with matter_env.factory() as db:
            await db.execute(delete(File).where(File.id == eid))
            await db.commit()


async def test_get_document_metadata_docx_returns_core_author(
    matter_env: MatterEnv, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A .docx's core-properties author / last-modified-by are read from the bytes."""
    from docx import Document as Docx

    d = Docx()
    d.add_paragraph("Body text.")
    d.core_properties.author = "Jane Author"
    d.core_properties.last_modified_by = "Mark Editor"
    buf = io.BytesIO()
    d.save(buf)
    data = buf.getvalue()

    @asynccontextmanager
    async def fake_download(*, storage_path: str) -> AsyncIterator[AsyncIterator[bytes]]:
        async def _gen() -> AsyncIterator[bytes]:
            yield data

        yield _gen()

    monkeypatch.setattr(agent_tools.storage, "stream_download", fake_download)

    async with matter_env.factory() as db:
        f = File(
            owner_id=matter_env.user_id,
            project_id=matter_env.project_id,
            filename="redline.docx",
            mime_type=OOXML_DOCX_MIME,
            size_bytes=len(data),
            hash_sha256="c" * 64,
            storage_path=f"meta-fixture/{uuid.uuid4()}",
            ingestion_status="ready",
        )
        db.add(f)
        await db.commit()
        fid = f.id
    try:
        out = await matter_env.get_meta("redline.docx")
        assert "DOCUMENT METADATA" in out
        assert "Jane Author" in out
        assert "Mark Editor" in out
        assert "forgeable" in out.lower()
    finally:
        async with matter_env.factory() as db:
            await db.execute(delete(File).where(File.id == fid))
            await db.commit()


async def test_get_document_metadata_unknown_name_lists_inventory(matter_env: MatterEnv) -> None:
    out = await matter_env.get_meta("ghost.eml")
    assert 'No document named "ghost.eml"' in out
    assert "msa.pdf" in out  # the inventory orients the model


# ---------------------------------------------------------------------------
# The guard chokepoint around every dispatch
# ---------------------------------------------------------------------------


async def test_each_dispatch_writes_one_audit_row_without_content(
    matter_env: MatterEnv,
) -> None:
    await matter_env.search("liability cap")
    await matter_env.read("msa.pdf")
    await matter_env.get_meta("msa.pdf")  # a pdf → "no structured metadata", still one row
    await matter_env.estimate(["msa.pdf"])  # F2 Slice E: read-cost estimate, one row

    rows = await _audit_rows(matter_env)
    assert len(rows) == 4
    assert {r.details["tool"] for r in rows} == MATTER_TOOL_NAMES
    for row in rows:
        assert row.action == "agent_run.tool_call"
        assert row.user_id == matter_env.user_id
        assert row.details["outcome"] == "success"
        assert isinstance(row.details["result_chars"], int)
        # Counts/types/IDs only — never document text or query text.
        serialized = str(row.details)
        assert "liability" not in serialized
        assert "aggregate" not in serialized


async def test_tools_expose_model_facing_schema(matter_env: MatterEnv) -> None:
    """LangChain builds the tool schema from name/docstring/signature —
    pin the model-visible surface (A-class content args only, ADR-F004)."""
    assert matter_env.search.__name__ == "search_documents"
    assert matter_env.read.__name__ == "read_document"
    assert matter_env.get_meta.__name__ == "get_document_metadata"
    assert matter_env.estimate.__name__ == "estimate_read_cost"
    assert matter_env.search.__doc__ and "empty query" in matter_env.search.__doc__
    assert matter_env.read.__doc__ and "filename" in matter_env.read.__doc__
    assert matter_env.get_meta.__doc__ and "forged" in matter_env.get_meta.__doc__
    assert matter_env.estimate.__doc__ and "estimate" in matter_env.estimate.__doc__.lower()
    assert list(inspect.signature(matter_env.search).parameters) == ["query"]
    assert list(inspect.signature(matter_env.read).parameters) == ["name"]
    assert list(inspect.signature(matter_env.get_meta).parameters) == ["name"]
    assert list(inspect.signature(matter_env.estimate).parameters) == ["filenames"]
    assert inspect.iscoroutinefunction(matter_env.search)
    assert inspect.iscoroutinefunction(matter_env.read)
    assert inspect.iscoroutinefunction(matter_env.get_meta)
    assert inspect.iscoroutinefunction(matter_env.estimate)


async def test_halted_run_denies_dispatch(matter_env: MatterEnv) -> None:
    """R5 at the tool boundary: a run no longer 'running' dispatches nothing."""
    from app.agents.guard import AgentRunHalted

    async with matter_env.factory() as db:
        run = await db.get(AgentRun, matter_env.run_id)
        assert run is not None
        run.status = "cancelled"
        await db.commit()

    with pytest.raises(AgentRunHalted):
        await matter_env.search("liability cap")

    rows = await _audit_rows(matter_env)
    assert rows[-1].details["outcome"] == "run_halted"


# ---------------------------------------------------------------------------
# End to end: the REAL deepagents loop dispatches the REAL guarded tool
# ---------------------------------------------------------------------------


async def test_real_loop_dispatches_guarded_search_over_matter_documents(
    matter_env: MatterEnv,
) -> None:
    """The slice's load-bearing integration: a scripted model drives the
    real loop, the model-initiated ``search_documents`` call runs the
    real FTS against the seeded matter, the result step carries the
    document name, and the dispatch left its audit row."""
    from app.agents.runner import execute_agent_run
    from app.models.agent_run import AgentRunStep
    from tests.agents.fakes import (
        ScriptedToolCallingModel,
        final_message,
        tool_call_message,
    )

    model = ScriptedToolCallingModel(
        responses=[
            tool_call_message("search_documents", {"query": "liability cap"}),
            final_message("The cap is the fees paid in the twelve months before the claim."),
        ]
    )

    await execute_agent_run(
        matter_env.run_id,
        matter_env.factory,
        tools=[matter_env.search, matter_env.read],
        model=model,
    )

    async with matter_env.factory() as db:
        run = (
            await db.execute(select(AgentRun).where(AgentRun.id == matter_env.run_id))
        ).scalar_one()
        steps = (
            (
                await db.execute(
                    select(AgentRunStep)
                    .where(AgentRunStep.run_id == matter_env.run_id)
                    .order_by(AgentRunStep.seq.asc())
                )
            )
            .scalars()
            .all()
        )

    assert run.status == "completed"
    assert run.final_answer is not None and "twelve" in run.final_answer
    kinds = [s.kind for s in steps]
    assert kinds == ["model_turn", "tool_call", "tool_result", "model_turn"]
    assert steps[1].name == "search_documents"
    assert "msa.pdf" in steps[2].summary  # the FTS hit, not canned text

    rows = await _audit_rows(matter_env)
    assert len(rows) == 1
    assert rows[0].details["tool"] == "search_documents"
    assert rows[0].details["outcome"] == "success"


# ---------------------------------------------------------------------------
# estimate_read_cost (F2 Slice E, ADR-F049): the pre-flight read-cost estimate
# ---------------------------------------------------------------------------


def test_read_tokens_estimate_and_cap() -> None:
    """Pure helper: ~chars/4, None→0, capped at the read limit (a truncated big
    doc cannot cost more than _READ_LIMIT/4 tokens to read)."""
    from app.agents.tools import _READ_LIMIT, _read_tokens

    assert _read_tokens(None) == 0
    assert _read_tokens(0) == 0
    assert _read_tokens(400) == 100
    assert _read_tokens(10**9) == _READ_LIMIT // 4  # capped at the read limit


async def test_estimate_read_cost_sums_named_documents(matter_env: MatterEnv) -> None:
    out = await matter_env.estimate(["msa.pdf", "notes.pdf"])
    expected = (len(_MSA_TEXT) + len(_NOTES_TEXT)) // 4
    assert "2 document(s)" in out
    assert f"~{expected:,} tokens" in out


async def test_estimate_read_cost_whole_matter_excludes_pending_and_decoys(
    matter_env: MatterEnv,
) -> None:
    """An empty list estimates the whole matter: the three READABLE matter docs
    (msa + notes + uploaded), never the pending file (no Document → not counted,
    zero cost), and never the other matter's or the foreign owner's documents."""
    out = await matter_env.estimate([])
    expected = (len(_MSA_TEXT) + len(_NOTES_TEXT) + len(_UPLOADED_TEXT)) // 4
    assert "the whole matter (3 document(s))" in out
    assert f"~{expected:,} tokens" in out


async def test_estimate_read_cost_unknown_name_is_honest(matter_env: MatterEnv) -> None:
    out = await matter_env.estimate(["does-not-exist.pdf"])
    assert "None of those filenames matched" in out


async def test_estimate_read_cost_excludes_other_matters_and_foreign_owner(
    matter_env: MatterEnv,
) -> None:
    """The matter+owner scope is the security boundary: a document in another matter
    of the same owner, or a foreign-owned file maliciously joined in, is invisible —
    estimate reports no match (404-conflated absence), never its cost."""
    assert "None of those filenames matched" in await matter_env.estimate(["secret.pdf"])
    assert "None of those filenames matched" in await matter_env.estimate(["foreign.pdf"])


async def test_estimate_read_cost_suggests_read_in_full_when_it_fits(
    matter_env: MatterEnv,
) -> None:
    out = await matter_env.estimate(["msa.pdf"])
    assert "read them in full" in out
    assert "estimate, not a live count" in out  # honest budget framing


async def test_estimate_read_cost_suggests_fan_out_when_over_budget(
    matter_env: MatterEnv,
) -> None:
    """A tiny context window makes even one small document exceed the remaining
    budget → the tool steers toward fan-out / passages, not read-in-full."""
    binding = MatterBinding(
        project_id=matter_env.project_id,
        user_id=matter_env.user_id,
        name="Acme MSA",
        privileged=False,
        minimum_inference_tier=None,
    )
    *_, estimate_tiny = build_matter_tools(
        matter_env.factory, run_id=matter_env.run_id, binding=binding, max_input_tokens=10_000
    )
    out = await estimate_tiny(["msa.pdf"])
    assert "too large to read whole" in out
    assert "fan out" in out.lower()
