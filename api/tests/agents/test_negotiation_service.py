"""C5a negotiation_service tests (ADR-F032).

The pure Adeu adapter: read the counterparty's tracked changes + comments into a
``StateOfPlay`` checklist, and apply per-ref decisions (accept/reject/counter/reply)
with the post-write reconciliation. Real Adeu (in-process, zero network) — no DB, no
gateway. This is the empirical floor under the no-silent-action guarantee: a decision
that doesn't land fails ``Reconciliation.ok``.
"""

from __future__ import annotations

import io
import zipfile

import pytest

from app.agents.negotiation_service import Decision, apply_decisions, read_state_of_play
from app.agents.redline_service import DEFAULT_AUTHOR

pytestmark = pytest.mark.integration


def _comment_authors_in_ooxml(docx_bytes: bytes) -> str:
    """Concatenated text of every ``word/comments*.xml`` part — the raw OOXML the live
    investigation inspected (a reply/thread that was wiped leaves no author here)."""
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        parts = [n for n in z.namelist() if n.startswith("word/comments") and n.endswith(".xml")]
        return "".join(z.read(n).decode("utf-8", "ignore") for n in parts)


def _base_docx() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_paragraph(
        "The Vendor shall indemnify the Customer for all losses arising from the Services."
    )
    doc.add_paragraph("The term of this Agreement is three (3) years from the Effective Date.")
    doc.add_paragraph("Liability is capped at fees paid in the prior twelve (12) months.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def counterparty_markup() -> bytes:
    """Three logical counterparty edits + 1 anchored comment (author='Opposing
    Counsel'). On adeu ≥2.x the two UNcommented edits each fan out into two
    word-level regions ("three"→"five" and "3"→"5"; the parens stay bare), so the
    document carries FIVE CriticMarkup regions — the state-of-play enumerates per
    REGION, exactly as it always has for real Word markup (which fragments
    arbitrarily too). Shared shape with the tool-integration tests."""
    from adeu import ModifyText, RedlineEngine

    eng = RedlineEngine(io.BytesIO(_base_docx()), author="Opposing Counsel")
    eng.apply_edits(
        [
            ModifyText(
                target_text="all losses",
                new_text="direct losses only",
                comment="Cap our exposure to direct losses please.",
            ),
            ModifyText(target_text="three (3)", new_text="five (5)"),
            ModifyText(target_text="twelve (12)", new_text="twenty-four (24)"),
        ]
    )
    out = eng.save_to_stream()
    return out.getvalue() if hasattr(out, "getvalue") else bytes(out)


def test_read_state_of_play_enumerates_changes_and_comments() -> None:
    state = read_state_of_play(counterparty_markup())

    # Document order, one ref per REGION (v2 surgical fan-out: 1 + 2 + 2).
    assert [c.ref for c in state.changes] == ["C1", "C2", "C3", "C4", "C5"]
    c1 = state.changes[0]
    assert c1.kind == "modify"
    assert c1.deleted_text == "all losses"  # commented edit stays one atomic region
    assert c1.inserted_text == "direct losses only"
    assert len(c1.adeu_ids) == 2  # a modify region is a del + ins pair
    assert [c.inserted_text for c in state.changes[1:]] == ["five", "5", "twenty-four", "24"]
    # the 2.x "(pairs with Chg:N)" meta suffix must never leak into the author
    assert all(c.author == "Opposing Counsel" for c in state.changes)

    # one OPEN counterparty comment (thread root, not ours)
    assert len(state.open_comment_refs) == 1
    cm = next(c for c in state.comments if c.parent_id is None and not c.is_ours)
    assert cm.author == "Opposing Counsel"
    assert "direct losses" in cm.text


def test_apply_decisions_full_round_reconciles() -> None:
    cp = counterparty_markup()
    state = read_state_of_play(cp)
    com_ref = next(iter(state.open_comment_refs))
    decisions = [
        # The comment is anchored to C1; accepting C1 deletes its thread, so the comment
        # is left_open (the accept is itself the answer) rather than replied to — replying
        # here would be wiped (covered by test_reply_then_reject_anchored_change_*).
        Decision(ref="C1", verdict="accept", rationale="Direct losses is acceptable."),
        # The term change is TWO regions on v2 ("three"→"five", "3"→"5") — reverting it
        # means a decision on each fragment (the coverage gate forces exactly this).
        Decision(ref="C2", verdict="reject", rationale="Five years is too long; revert to three."),
        Decision(ref="C3", verdict="reject", rationale="Five years is too long; revert to three."),
        # The cap change is also two regions; the counter rides the word fragment and
        # rewrites the whole phrase, the digit fragment is accepted (the counter's own
        # word-diff then strikes the accepted "24" → "18").
        Decision(
            ref="C4",
            verdict="counter",
            target_text="twenty-four (24)",
            new_text="eighteen (18)",
            rationale="Eighteen months is the house fallback for the liability survival window here.",
        ),
        Decision(ref="C5", verdict="accept", rationale="Digits follow the countered wording."),
        Decision(ref=com_ref, verdict="leave_open", rationale="Accepting their edit answers this."),
    ]

    out_bytes, recon = apply_decisions(cp, state, decisions)

    assert recon.ok, recon.issues
    assert recon.review_skipped == 0
    # Four modify pairs got a same-verdict decision (C1 accept, C2+C3 reject,
    # C5 accept); each pair's second id counts already_resolved on adeu >=2.x —
    # pins the 3-tuple unpack ORDER, not just its arity.
    assert recon.review_already_resolved == 4
    assert recon.counters_skipped == 0
    assert recon.counters_applied == 1  # per LOGICAL counter on v2, not per region

    after = read_state_of_play(out_bytes)
    # C1 accepted → incorporated as clean text; C2+C3 rejected → back to "three (3)".
    assert "direct losses only" in after.clean_view
    assert "three (3)" in after.clean_view
    assert "five" not in after.clean_view
    # the counter is a pending tracked change: accept-all view shows our number
    assert "eighteen (18)" in after.clean_view


def test_read_state_of_play_captures_comment_anchor() -> None:
    """C5b-1: the comment co-occurs in C1's meta block, so it anchors to C1 (not C2)."""
    state = read_state_of_play(counterparty_markup())
    com_ref = next(iter(state.open_comment_refs))
    assert state.comment_anchors == {com_ref: "C1"}
    # anchors always point at a real logical change ref, never a raw Chg id
    assert set(state.comment_anchors.values()) <= state.change_refs


def test_reply_then_reject_anchored_change_fails_reconciliation() -> None:
    """The comment-wipe bug, caught at the document level: rejecting the change the
    comment is anchored to deletes the thread (and our reply) — Adeu still reports it
    applied, so only the reply-survival reconciliation catches it. Nothing is persisted."""
    cp = counterparty_markup()
    state = read_state_of_play(cp)
    com_ref = next(iter(state.open_comment_refs))  # anchored to C1
    decisions = [
        Decision(ref="C1", verdict="reject", rationale="Keep the broader indemnity."),
        Decision(ref=com_ref, verdict="reply", reply_text="We need full indemnity cover here."),
    ]

    out_bytes, recon = apply_decisions(cp, state, decisions)

    assert not recon.ok
    assert any("survive" in issue for issue in recon.issues)
    # OOXML-level proof (the same inspection that found the live bug): our reply is gone.
    assert DEFAULT_AUTHOR not in _comment_authors_in_ooxml(out_bytes)


def test_reply_survives_when_anchored_change_untouched() -> None:
    """The safe path: reply to the anchored comment while NOT accepting/rejecting its
    change (here we accept a DIFFERENT change) → the reply survives, reconciliation ok."""
    cp = counterparty_markup()
    state = read_state_of_play(cp)
    com_ref = next(iter(state.open_comment_refs))  # anchored to C1
    decisions = [
        Decision(ref="C2", verdict="accept", rationale="Five years is fine."),
        Decision(ref=com_ref, verdict="reply", reply_text="Agreed — see our position."),
    ]

    out_bytes, recon = apply_decisions(cp, state, decisions)

    assert recon.ok, recon.issues
    assert DEFAULT_AUTHOR in _comment_authors_in_ooxml(out_bytes)  # our reply is present
    after = read_state_of_play(out_bytes)
    assert any(cm.is_ours and cm.parent_id is not None for cm in after.comments)


def test_apply_decisions_unknown_ref_fails_reconciliation() -> None:
    cp = counterparty_markup()
    state = read_state_of_play(cp)
    _out, recon = apply_decisions(cp, state, [Decision(ref="C99", verdict="accept", rationale="x")])
    assert not recon.ok
    assert any("C99" in issue for issue in recon.issues)


def test_apply_decisions_bad_anchor_counter_is_skipped_not_silent() -> None:
    """A counter whose target_text isn't in the document can't land → reconciliation
    fails (the response is rejected upstream; nothing is silently dropped)."""
    cp = counterparty_markup()
    state = read_state_of_play(cp)
    decisions = [
        Decision(
            ref="C1",
            verdict="counter",
            target_text="this exact phrase is absent from the document",
            new_text="replacement that will never anchor",
            rationale="A deliberately unanchorable counter to prove the reconciliation tripwire.",
        )
    ]
    _out, recon = apply_decisions(cp, state, decisions)
    assert not recon.ok


def test_pairs_suffix_regex_strips_both_forms_and_spares_real_parentheticals() -> None:
    """Direct drift-guard for the 2.x meta-line author suffix (both documented
    shapes), without over-stripping a genuine parenthetical in an author name."""
    from app.agents.negotiation_service import _PAIRS_SUFFIX

    cases = {
        "Jane Smith (pairs with Chg:6)": "Jane Smith",
        "Opposing Counsel (pairs Chg:2, Chg:3)": "Opposing Counsel",
        "Ltd (UK) counsel (pairs with Chg:9)": "Ltd (UK) counsel",
        "Plain Author": "Plain Author",
        "Acme (Holdings) Ltd": "Acme (Holdings) Ltd",
    }
    for raw, want in cases.items():
        assert _PAIRS_SUFFIX.sub("", raw).strip() == want, raw
