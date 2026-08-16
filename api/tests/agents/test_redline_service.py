"""C4 RedlineService + import-boundary tests (ADR-F031).

Exercises the Adeu SDK adapter on the pin (validate → two-pass render → preview
→ apply → accept-all; the ADR-F085 recipe) and enforces the STRICT import
boundary: app code must never import ``adeu.server`` / ``adeu.mcp_components``
(a second network egress) — the C4 analogue of the C1 fitz import-guard.
"""

from __future__ import annotations

import ast
import io
import pathlib
import re
import zipfile

from app.agents.redline_render import reconstruct_redline_text
from app.agents.redline_service import (
    ProposedEdit,
    RedlineService,
    ensure_track_changes_recording,
)

CAP = (
    "The Vendor's aggregate liability arising out of or in connection with this "
    "Agreement shall not exceed the total fees paid by the Customer in the three "
    "(3) months preceding the claim."
)


def _build_docx(paragraphs: list[str]) -> bytes:
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_text(data: bytes) -> str:
    from docx import Document

    return "\n".join(p.text for p in Document(io.BytesIO(data)).paragraphs)


def _strip_markers(redline: str) -> str:
    """Drop tracked-change spans, leaving only the unchanged (bare) text."""
    no_ins = re.sub(r"\[\+.*?\+\]", "", redline, flags=re.DOTALL)
    return re.sub(r"\[-.*?-\]", "", no_ins, flags=re.DOTALL)


def _settings_xml(data: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        return (
            z.read("word/settings.xml").decode("utf-8")
            if "word/settings.xml" in z.namelist()
            else ""
        )


# --------------------------------------------------------------------------- #
# Track-changes recording — ADR-F047 Slice 5 (editor hand-back)
# --------------------------------------------------------------------------- #


def test_apply_output_opens_with_record_changes_on() -> None:
    """A redline must open with "Record Changes" ON so the lawyer's edits in the editor
    are captured as tracked changes (Adeu emits tracked content but not the flag)."""
    from app.agents.negotiation_service import read_state_of_play

    src = _build_docx([CAP])
    out = RedlineService().apply(
        src, [ProposedEdit(target_text="three (3)", new_text="twelve (12)")]
    )
    assert "w:trackChanges" in _settings_xml(out.docx_bytes)  # recording on
    # the agent's tracked content is intact + re-readable (and authored by the agent)
    state = read_state_of_play(out.docx_bytes)
    assert len(state.changes) >= 1


def test_ensure_track_changes_recording_is_idempotent() -> None:
    src = _build_docx([CAP])
    once = ensure_track_changes_recording(src)
    assert "w:trackChanges" in _settings_xml(once)
    twice = ensure_track_changes_recording(once)
    assert twice == once  # already recording → byte-identical no-op


def test_ensure_track_changes_recording_preserves_other_parts() -> None:
    """Surgical: only settings.xml changes; document.xml stays byte-identical."""
    src = _build_docx([CAP, "Second paragraph."])
    patched = ensure_track_changes_recording(src)
    with zipfile.ZipFile(io.BytesIO(src)) as a, zipfile.ZipFile(io.BytesIO(patched)) as b:
        assert a.read("word/document.xml") == b.read("word/document.xml")
    assert _docx_text(patched) == _docx_text(src)


def _with_settings_track_changes_off(data: bytes) -> bytes:
    """Inject Word's explicit "tracking OFF" flag (`<w:trackChanges w:val="false"/>`)."""
    with zipfile.ZipFile(io.BytesIO(data)) as zin:
        entries = [(n, zin.read(n)) for n in zin.namelist()]
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for n, d in entries:
            if n == "word/settings.xml":
                s = re.sub(
                    r"(<w:settings\b[^>]*?>)",
                    r'\1<w:trackChanges w:val="false"/>',
                    d.decode(),
                    count=1,
                )
                d = s.encode()
            zout.writestr(n, d)
    return out.getvalue()


def test_ensure_track_changes_recording_forces_on_when_explicitly_off() -> None:
    """A source the lawyer saved with Track Changes OFF carries `w:val="false"` — the
    editor would open with recording off. ensure_track_changes_recording must flip it ON,
    not be fooled by the substring into a no-op (the C-review should-fix)."""
    src = _with_settings_track_changes_off(_build_docx([CAP]))
    assert 'w:val="false"' in _settings_xml(src)  # precondition: explicitly OFF
    on = _settings_xml(ensure_track_changes_recording(src))
    assert "trackChanges" in on
    assert 'w:val="false"' not in on and 'w:val="0"' not in on  # recording forced ON


def test_ensure_track_changes_recording_is_schema_ordered_after_zoom() -> None:
    """CT_Settings is an ordered sequence; the recording flag must be inserted AFTER the
    early elements (python-docx's default settings.xml leads with <w:zoom>), not as the
    first child (the C-review should-fix on schema order)."""
    s = _settings_xml(ensure_track_changes_recording(_build_docx([CAP])))
    assert "zoom" in s and "trackChanges" in s
    assert s.index("zoom") < s.index("trackChanges")


# --------------------------------------------------------------------------- #
# Import boundary — the load-bearing security guard
# --------------------------------------------------------------------------- #


def _dotted(node: ast.AST) -> str | None:
    parts: list[str] = []
    while isinstance(node, ast.Attribute):
        parts.append(node.attr)
        node = node.value
    if isinstance(node, ast.Name):
        parts.append(node.id)
        return ".".join(reversed(parts))
    return None


def test_app_never_imports_adeu_server() -> None:
    """No app module may import or reference Adeu's bundled server surfaces
    (a second egress). Adeu is SDK-only (ADR-F031). ``adeu.serve`` (2.x's
    JSON-lines daemon — one letter off ``adeu.server``) and ``adeu.cli`` both
    transitively import ``adeu.mcp_components``, so they are banned by name too
    (the F085 review found the one-letter gap)."""
    app_dir = pathlib.Path(__file__).resolve().parents[1].parent / "app"
    banned = ("adeu.server", "adeu.serve", "adeu.cli", "adeu.mcp_components")
    offenders: list[tuple[str, str]] = []
    for py in sorted(app_dir.rglob("*.py")):
        tree = ast.parse(py.read_text(encoding="utf-8"))
        for node in ast.walk(tree):
            if isinstance(node, ast.Import):
                for alias in node.names:
                    if any(alias.name == b or alias.name.startswith(b + ".") for b in banned):
                        offenders.append((py.name, alias.name))
            elif isinstance(node, ast.ImportFrom) and node.module:
                if any(node.module == b or node.module.startswith(b + ".") for b in banned):
                    offenders.append((py.name, node.module))
            elif isinstance(node, ast.Attribute):
                dotted = _dotted(node)
                if dotted and any(dotted == b or dotted.startswith(b + ".") for b in banned):
                    offenders.append((py.name, dotted))
    assert offenders == [], (
        "app code must not import/reference adeu.server / adeu.mcp_components — Adeu "
        f"is SDK-only (ADR-F031). Offenders: {offenders}"
    )


# --------------------------------------------------------------------------- #
# SDK adapter behaviour (needs adeu + python-docx)
# --------------------------------------------------------------------------- #


def test_raw_edit_is_surgical_and_does_not_corrupt() -> None:
    """A multi-token replace marks only the changed words; the rest of the clause
    stays bare and uncorrupted (the regression guard against the historical
    micro-anchor corruption, e.g. 'Ven12or' — ADR-F045 word-diff renders
    positionally, not by fuzzy micro-match)."""
    svc = RedlineService()
    docx = _build_docx([CAP])
    res = svc.apply(docx, [ProposedEdit("three (3) months", "twelve (12) months", None)])
    redline = reconstruct_redline_text(res.docx_bytes)
    assert "The Vendor's aggregate liability arising out of or in connection" in redline
    assert "Ven12or" not in redline and "Ven[" not in redline
    clean = _docx_text(svc.accept_all(res.docx_bytes))
    assert clean.startswith("The Vendor's aggregate liability")
    assert "twelve (12) months" in clean and "three (3) months" not in clean


def test_carveout_append_renders_surgically() -> None:
    svc = RedlineService()
    docx = _build_docx([CAP])
    res = svc.apply(
        docx,
        [
            ProposedEdit(
                "preceding the claim.",
                "preceding the claim, save that data-protection liability shall be unlimited.",
                "carve data protection out of the cap",
            )
        ],
    )
    redline = reconstruct_redline_text(res.docx_bytes)
    assert "[+" in redline  # protective language inserted
    # the clause body stays bare; only the boundary token + addition are tracked
    assert "Customer in the three (3) months preceding the" in _strip_markers(redline)
    clean = _docx_text(svc.accept_all(res.docx_bytes))
    assert "save that data-protection liability shall be unlimited" in clean


def test_preview_apply_and_accept_roundtrip() -> None:
    svc = RedlineService()
    docx = _build_docx([CAP])
    edits = [ProposedEdit("three (3) months", "twelve (12) months", "align to house floor")]

    preview = svc.dry_run(docx, edits)
    assert preview.edits_applied >= 1
    assert preview.edits_skipped == 0

    result = svc.apply(docx, edits)
    redline = reconstruct_redline_text(result.docx_bytes)
    assert "[+twelve" in redline and "[-three" in redline  # native tracked changes
    # unchanged head stays bare (surgical)
    assert "shall not exceed the total fees paid by the Customer" in redline

    clean = _docx_text(svc.accept_all(result.docx_bytes))
    assert "twelve (12) months" in clean
    assert "three (3) months" not in clean


def test_apply_skips_unanchored_edit() -> None:
    """An edit whose target isn't in the document is reported skipped, not applied."""
    svc = RedlineService()
    docx = _build_docx([CAP])
    preview = svc.dry_run(docx, [ProposedEdit("this phrase is absent here", "x y z")])
    assert preview.edits_skipped >= 1


# --------------------------------------------------------------------------- #
# Word-level diff rendering (ADR-F045) — the TOOL keeps unchanged wording bare,
# even when the model quotes a whole clause as one edit.
# --------------------------------------------------------------------------- #

# A realistic multi-paragraph contract: the indemnity is NOT at offset 0, and
# "Customer"/"Vendor" recur — so a positional misplacement would be visible.
_MSA = [
    "MASTER SERVICES AGREEMENT",
    '1. Definitions. "Customer" means the entity identified above; "Vendor" means the supplier.',
    "2. Services. The Vendor shall provide the Services to the Customer in accordance with the Order.",
    (
        "8. Indemnity. The Customer shall indemnify, defend and hold harmless the Vendor and its "
        "affiliates against any and all claims, losses, damages, liabilities and expenses arising "
        "from or in connection with the Customer use of the Services or the Customer Data."
    ),
    "9. Fees. The Customer shall pay the Vendor within thirty (30) days.",
    "10. Term. This Agreement runs for one (1) year from the Effective Date.",
]
_INDEMNITY = _MSA[3].split("Indemnity. ", 1)[1]  # the clause text after the heading
_INDEMNITY_MUTUAL = (
    "Each party shall indemnify, defend and hold harmless the other party and its "
    "affiliates against any and all claims, losses, damages, liabilities and expenses arising "
    "from or in connection with a party breach of this Agreement."
)


def _bare(redline: str) -> str:
    """The untouched (neither struck nor inserted) text of a reconstruction."""
    return _strip_markers(redline)


def test_worddiff_keeps_clause_interior_bare() -> None:
    """A whole-clause mutualisation quoted as ONE edit renders as several minimal
    regions — the indemnity verb phrase stays bare, every other paragraph intact
    (the C8/C9 swallow fix)."""
    svc = RedlineService()
    docx = _build_docx(_MSA)
    res = svc.apply(
        docx, [ProposedEdit(_INDEMNITY, _INDEMNITY_MUTUAL, "Mutualise indemnity; narrow trigger.")]
    )
    redline = reconstruct_redline_text(res.docx_bytes)
    bare = _bare(redline)

    # the recognisable boilerplate is never touched …
    assert "shall indemnify, defend and hold harmless" in bare
    assert "any and all claims, losses, damages, liabilities and expenses" in bare
    # … only the party/trigger words are struck (several regions, not one block) …
    assert redline.count("[-") >= 3
    assert "[+other party+]" in redline  # protected party mutualised
    assert re.search(r"\[\+Each", redline)  # indemnifying party mutualised
    # … no micro-anchor corruption …
    assert not re.search(r"[A-Za-z]\d[A-Za-z]", redline)
    # … and the untouched paragraphs survive verbatim.
    assert "The Vendor shall provide the Services to the Customer" in bare
    assert "within thirty (30) days" in bare
    assert "one (1) year from the Effective Date" in bare

    clean = _docx_text(svc.accept_all(res.docx_bytes))
    assert "Each party shall indemnify, defend and hold harmless the other party" in clean
    assert "The Customer shall indemnify" not in clean


def test_worddiff_multi_edit_batch_no_cross_contamination() -> None:
    """Two whole-clause edits in one batch each render surgically and land in their
    own clause (positional, full-document coordinates)."""
    svc = RedlineService()
    docx = _build_docx(_MSA)
    res = svc.apply(
        docx,
        [
            ProposedEdit(_INDEMNITY, _INDEMNITY_MUTUAL, "Mutualise indemnity."),
            ProposedEdit(
                "The Customer shall pay the Vendor within thirty (30) days.",
                "The Customer shall pay the Vendor within sixty (60) days.",
                "Extend the payment period to 60 days.",
            ),
        ],
    )
    redline = reconstruct_redline_text(res.docx_bytes)
    bare = _bare(redline)
    assert "shall indemnify, defend and hold harmless" in bare  # boilerplate bare
    assert "shall pay the Vendor within" in bare  # fee-clause stem bare
    assert "[-thirty-]" in redline and "[+sixty+]" in redline
    assert "one (1) year from the Effective Date" in bare  # term clause untouched

    clean = _docx_text(svc.accept_all(res.docx_bytes))
    assert "within sixty (60) days" in clean
    assert "Each party shall indemnify" in clean


def test_worddiff_genuine_rewrite_still_renders_as_block() -> None:
    """Word-diff does NOT fabricate surgery: a true rewrite (every word changed)
    legitimately renders as one struck-and-retyped block — so the gate, not the
    renderer, is what guards genuine over-rewording."""
    svc = RedlineService()
    docx = _build_docx(_MSA)
    res = svc.apply(
        docx,
        [
            ProposedEdit(
                "The Vendor shall provide the Services to the Customer in accordance with the Order.",
                "Supplier will deliver outputs per each statement of work executed hereunder.",
                "Full rewrite of the services clause.",
            )
        ],
    )
    redline = reconstruct_redline_text(res.docx_bytes)
    # one contiguous struck region + one inserted region (a real rewrite)
    assert redline.count("[-") == 1 and redline.count("[+") == 1
    assert not re.search(r"[A-Za-z]\d[A-Za-z]", redline)  # still no corruption


def test_worddiff_hyphenated_terms_do_not_corrupt() -> None:
    """Hyphenated/compound replacements render cleanly (the historical
    '-'/'_' mid-word-split corruption does not reproduce on the pin)."""
    svc = RedlineService()
    docx = _build_docx(_MSA)
    res = svc.apply(
        docx,
        [
            ProposedEdit(
                '"Vendor" means the supplier.',
                '"Supplier" means the service-provider.',
                "Rename the defined term and use a hyphenated description.",
            )
        ],
    )
    redline = reconstruct_redline_text(res.docx_bytes)
    assert not re.search(r"[A-Za-z]\d[A-Za-z]", redline)
    clean = _docx_text(svc.accept_all(res.docx_bytes))
    assert "the service-provider" in clean


def test_ambiguous_anchor_rejected_with_actionable_problem() -> None:
    """A non-unique anchor is REJECTED with a per-edit message (Adeu 2.x's
    validate_edits shows the competing contexts) instead of the pre-2.x shim's
    silent wholesale fallback onto the first occurrence — the gate's D4 forbids
    ambiguity anyway, so the service must never guess a location."""
    svc = RedlineService()
    # "The Customer" appears in several paragraphs → ambiguous under strict mode.
    docx = _build_docx(_MSA)
    preview = svc.dry_run(docx, [ProposedEdit("The Customer", "Each party")])
    assert preview.edits_applied == 0
    assert preview.edits_skipped >= 1
    assert preview.problems  # the model gets told WHICH edit and WHY


# --------------------------------------------------------------------------- #
# Two-pass recipe (ADR-F085) — comments span the logical edit; living redline
# round 2 resolves over prior tracked changes (the probe scenarios as tests:
# docs/fork/evidence/adeu2-probe/).
# --------------------------------------------------------------------------- #


def _ooxml_regions(docx_bytes: bytes) -> list[str]:
    """local tag names of every tracked-change region in document order."""
    import re as _re

    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        xml = z.read("word/document.xml").decode("utf-8", "ignore")
    return _re.findall(r"<w:(ins|del)\b", xml)


def _comments_xml(docx_bytes: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        parts = [n for n in z.namelist() if n.startswith("word/comments") and n.endswith(".xml")]
        return "".join(z.read(n).decode("utf-8", "ignore") for n in parts)


def test_commented_edit_is_surgical_and_comment_spans_the_clause() -> None:
    """The recipe's whole point: a comment-carrying edit still fans out into
    minimal regions (the uncommented apply pass), and ONE rationale comment
    anchors across the clause (the comment-only annotate pass) — not Adeu's atomic
    one-block commented rendering, and not the old comment-on-first-fragment."""
    svc = RedlineService()
    docx = _build_docx(_MSA)
    res = svc.apply(
        docx, [ProposedEdit(_INDEMNITY, _INDEMNITY_MUTUAL, "Mutualise indemnity; narrow trigger.")]
    )
    assert res.comments_applied == 1
    # several minimal regions, not one struck-and-retyped block
    assert len(_ooxml_regions(res.docx_bytes)) >= 4
    comments = _comments_xml(res.docx_bytes)
    assert "Mutualise indemnity" in comments
    # the comment is range-anchored in the body (rejecting one word-region in
    # review cannot orphan it)
    with zipfile.ZipFile(io.BytesIO(res.docx_bytes)) as z:
        body = z.read("word/document.xml").decode("utf-8", "ignore")
    assert "commentRangeStart" in body and "commentRangeEnd" in body


def test_living_redline_round2_resolves_over_round1_insertion() -> None:
    """Round 2 on the redlined output (the ADR-F081 living redline): a target that
    INCLUDES text round 1 inserted must resolve and land in the right clause, with
    its rationale comment attached — the case where both the pre-2.x pinned shim
    (silent wrong-clause mis-anchor, AP-05) and a naive commented native apply
    (silent comment drop) fail."""
    svc = RedlineService()
    docx = _build_docx(_MSA)
    r1 = svc.apply(docx, [ProposedEdit(_INDEMNITY, _INDEMNITY_MUTUAL, "Mutualise indemnity.")])
    # Round 2 targets a span that only exists because round 1 inserted it.
    r2 = svc.apply(
        r1.docx_bytes,
        [
            ProposedEdit(
                "a party breach of this Agreement.",
                "a party material breach of this Agreement.",
                "Limit the indemnity trigger to material breach.",
            )
        ],
    )
    assert r2.edits_applied == 1 and r2.edits_skipped == 0
    assert r2.comments_applied == 1
    comments = _comments_xml(r2.docx_bytes)
    assert "Mutualise indemnity" in comments and "material breach" in comments
    clean = _docx_text(svc.accept_all(r2.docx_bytes))
    assert "a party material breach of this Agreement" in clean
    # no mis-anchor: the untouched clauses survive verbatim
    assert "within thirty (30) days" in clean
    assert "one (1) year from the Effective Date" in clean


def test_round2_target_inside_deleted_text_is_rejected() -> None:
    """Text round 1 struck (pending deletion) is not editable — the edit must come
    back as a per-edit problem, never land inside the deletion."""
    svc = RedlineService()
    docx = _build_docx(_MSA)
    r1 = svc.apply(docx, [ProposedEdit(_INDEMNITY, _INDEMNITY_MUTUAL, "Mutualise.")])
    preview = svc.dry_run(
        r1.docx_bytes,
        # round 1 deleted "the Customer use of the Services" (trigger rewritten)
        [ProposedEdit("the Customer use of the Services", "the Customer's use")],
    )
    assert preview.edits_applied == 0
    assert preview.edits_skipped >= 1
    assert preview.problems


def test_dry_run_problems_name_the_failing_edit() -> None:
    """Per-edit attribution (VM2-A): with a clean first edit and a bad second one,
    the problem message names edit 2 — the batch index space is the model's own
    edit list."""
    svc = RedlineService()
    docx = _build_docx(_MSA)
    preview = svc.dry_run(
        docx,
        [
            ProposedEdit("within thirty (30) days.", "within sixty (60) days."),
            ProposedEdit("this phrase is absent from the document", "x"),
        ],
    )
    assert preview.edits_applied == 0  # all-or-nothing batch
    assert any("Edit 2" in p for p in preview.problems)


def test_repeated_new_text_still_carries_its_comment() -> None:
    """The ordering discriminator (ADR-F085): the rationale comment anchors on
    target_text (D4-unique), NOT on new_text — so an edit whose new_text already
    appears verbatim elsewhere in the document still gets its comment. Under the
    rejected apply-first ordering (comment-only on new_text after the apply) this
    exact case skips the comment as ambiguous and the batch rejects."""
    svc = RedlineService()
    docx = _build_docx(
        [
            "1. Invoices under a Statement of Work are payable within sixty (60) days.",
            "2. All other invoices are payable within thirty (30) days.",
        ]
    )
    res = svc.apply(
        docx,
        [
            ProposedEdit(
                "All other invoices are payable within thirty (30) days.",
                # new_text's payment phrase already exists verbatim in clause 1
                "All other invoices are payable within sixty (60) days.",
                "Align the general payment period with the SOW period per house terms.",
            )
        ],
    )
    assert res.edits_applied == 1
    assert res.comments_applied == 1  # anchored on the unique target, never ambiguous
    assert "house terms" in _comments_xml(res.docx_bytes)
    clean = _docx_text(svc.accept_all(res.docx_bytes))
    assert clean.count("sixty (60)") == 2 and "thirty (30)" not in clean
