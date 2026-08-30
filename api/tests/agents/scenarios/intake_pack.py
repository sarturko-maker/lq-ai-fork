"""Rig for the INTAKE-3 intake-run evals — seed a mailbox, land an envelope, run it.

Shared by ``test_intake_outcome_eval.py`` (the code-scored outcome gate) and
``test_intake_run_live.py`` (the two named live scenarios). The point of the rig is
that it drives the PRODUCTION path end to end:

``POST /internal/intake/emails``'s own handler (:func:`app.api.intake_emails.ingest_email`)
lands the envelope → the real worker core
(:func:`app.workers.intake_worker.process_intake_thread`) builds the intake prompt and
creates the run → the real composition point
(:func:`app.agents.composition.compose_and_execute_run`) executes it against the live
gateway with Commercial's bound skills.

Two seams are injected, both for the test process only:

* ``enqueue`` — the worker's arq handoff is replaced by a no-op that reports success
  (the run is executed inline here, not by the dev worker);
* a ``checkpointer`` — an intake run ALWAYS compiles a HITL floor
  (``draft_email_reply`` is granted and interrupt-gated structurally), and the runner
  fails closed without one. Production runs on the Postgres checkpointer; these tests
  use an in-memory saver.

Attachments are ingested for real (storage + a ``files`` row), then their text is
materialised into ``documents`` here — the ingest WORKER is a separate container
against a different database, so without this the agent would only ever see
"ingestion pending" and could never read an attached NDA.
"""

from __future__ import annotations

import base64
import io
import json
import uuid
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from langgraph.checkpoint.memory import InMemorySaver
from sqlalchemy import delete, select
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker

from app.agents.composition import compose_and_execute_run
from app.agents.intake_tools import safe_fail_intake_thread
from app.api.intake_emails import ingest_email
from app.config import get_settings
from app.models.agent_run import AgentRun, AgentRunStep, AgentThread
from app.models.audit import AuditLog
from app.models.document import Document, DocumentChunk
from app.models.file import File
from app.models.intake import IntakeMailbox, IntakeMessage, IntakeThread
from app.models.practice_area import PracticeArea
from app.models.project import Project
from app.models.user import User
from app.schemas.intake import InboundEmailEnvelope
from app.security import hash_password
from app.skills import load_registry
from app.workers.intake_worker import process_intake_thread

PACK_DIR = Path(__file__).resolve().parents[4] / "sample-documents" / "commercial-intake-pack"
SKILLS_DIR = Path("/skills")

# The pack's inbox — every envelope names it, so ONE mailbox binding serves them all.
PACK_INBOX_ID = "legal-intake@northwindtrading.co.uk"


@dataclass(frozen=True)
class Fixture:
    """One committed envelope plus its code-scored expectation."""

    name: str
    envelope: dict[str, Any]
    expected: str
    # The outcome that would be a SAFETY failure on this thread. ``None`` where none
    # exists: with two outcomes (ADR-F086 A1) over-caution on noise is never unsafe.
    unsafe_if: str | None


@dataclass
class IntakeRig:
    factory: async_sessionmaker[AsyncSession]
    user_id: uuid.UUID
    mailbox_id: uuid.UUID


@dataclass
class IntakeResult:
    """What one intake run left behind — the structural record, not model prose."""

    fixture: str
    thread_id: uuid.UUID
    project_id: uuid.UUID | None
    run_id: uuid.UUID | None
    run_status: str | None
    outcome: str | None
    label: str | None
    thread_status: str
    project_archived: bool
    tools_called: list[str]
    error: str | None


def load_pack(only: list[str] | None = None) -> list[Fixture]:
    """Every fixture in ``sample-documents/commercial-intake-pack``, in file order."""
    expected = json.loads((PACK_DIR / "expected.json").read_text())
    fixtures: list[Fixture] = []
    for name in sorted(expected):
        if only is not None and name not in only:
            continue
        envelope = json.loads((PACK_DIR / name).read_text())
        fixtures.append(
            Fixture(
                name=name,
                envelope=envelope,
                expected=expected[name]["expected"],
                unsafe_if=expected[name]["unsafe_if"],
            )
        )
    return fixtures


async def seed_rig(factory: async_sessionmaker[AsyncSession]) -> IntakeRig:
    """One queue-owner user + one Commercial-bound mailbox for the pack's inbox."""
    async with factory() as db:
        area_id = (
            await db.execute(select(PracticeArea.id).where(PracticeArea.key == "commercial"))
        ).scalar_one()
        user = User(
            email=f"intake-eval-{uuid.uuid4().hex[:8]}@example.com",
            display_name="Intake Queue Owner",
            hashed_password=hash_password("correct-horse-battery-staple"),
            is_admin=False,
            mfa_enabled=False,
            must_change_password=False,
        )
        db.add(user)
        await db.flush()
        mailbox = IntakeMailbox(
            provider="agentmail",
            inbox_id=PACK_INBOX_ID,
            address=PACK_INBOX_ID,
            practice_area_id=area_id,
            owner_user_id=user.id,
        )
        db.add(mailbox)
        await db.commit()
        return IntakeRig(factory=factory, user_id=user.id, mailbox_id=mailbox.id)


async def teardown_rig(rig: IntakeRig) -> None:
    async with rig.factory() as db:
        await db.execute(delete(AuditLog).where(AuditLog.user_id == rig.user_id))
        run_ids = (
            (await db.execute(select(AgentRun.id).where(AgentRun.user_id == rig.user_id)))
            .scalars()
            .all()
        )
        if run_ids:
            await db.execute(delete(AgentRunStep).where(AgentRunStep.run_id.in_(run_ids)))
        # intake_threads/messages CASCADE with the mailbox.
        await db.execute(delete(IntakeMailbox).where(IntakeMailbox.id == rig.mailbox_id))
        await db.execute(delete(AgentRun).where(AgentRun.user_id == rig.user_id))
        await db.execute(delete(AgentThread).where(AgentThread.user_id == rig.user_id))
        # documents/document_chunks are FK'd to files, not to the user — delete them
        # explicitly rather than relying on a cascade this rig does not own.
        file_ids = (
            (await db.execute(select(File.id).where(File.owner_id == rig.user_id))).scalars().all()
        )
        if file_ids:
            doc_ids = (
                (await db.execute(select(Document.id).where(Document.file_id.in_(file_ids))))
                .scalars()
                .all()
            )
            if doc_ids:
                await db.execute(
                    delete(DocumentChunk).where(DocumentChunk.document_id.in_(doc_ids))
                )
            await db.execute(delete(Document).where(Document.file_id.in_(file_ids)))
        await db.execute(delete(File).where(File.owner_id == rig.user_id))
        await db.execute(delete(Project).where(Project.owner_id == rig.user_id))
        await db.execute(delete(User).where(User.id == rig.user_id))
        await db.commit()


async def _always_enqueued(_run_id: uuid.UUID) -> bool:
    """The run is executed inline by this rig, so the arq handoff is a no-op."""
    return True


async def land(rig: IntakeRig, fixture: Fixture) -> tuple[uuid.UUID, uuid.UUID | None]:
    """Land one envelope through the REAL bridge-landing handler."""
    envelope = InboundEmailEnvelope.model_validate(fixture.envelope)
    async with rig.factory() as db:
        response = await ingest_email(envelope, get_settings(), db)
    assert response.thread_id is not None, f"{fixture.name}: landing produced no thread"
    return response.thread_id, response.project_id


async def materialise_attachment_text(rig: IntakeRig, project_id: uuid.UUID) -> int:
    """Give every ingested .docx a ``documents`` row so ``read_document`` works.

    The ingest worker is a separate container pointed at the dev database; in a test
    process nothing else turns the stored bytes into extractable text.
    """
    written = 0
    async with rig.factory() as db:
        rows = (
            (
                await db.execute(
                    select(File).where(File.project_id == project_id, File.deleted_at.is_(None))
                )
            )
            .scalars()
            .all()
        )
        for file in rows:
            existing = (
                await db.execute(select(Document.id).where(Document.file_id == file.id))
            ).scalar_one_or_none()
            if existing is not None or not file.filename.lower().endswith(".docx"):
                continue
            data = _ATTACHMENT_BYTES.get(file.filename)
            if data is None:
                continue
            text = _docx_text(data)
            document = Document(
                file_id=file.id,
                parser="intake-eval-fixture",
                page_count=1,
                character_count=len(text),
                normalized_content=text,
            )
            db.add(document)
            await db.flush()
            db.add(
                DocumentChunk(
                    document_id=document.id,
                    chunk_index=0,
                    content=text,
                    page_start=1,
                    page_end=1,
                    char_offset_start=0,
                    char_offset_end=len(text),
                )
            )
            file.ingestion_status = "ready"
            written += 1
        await db.commit()
    return written


_ATTACHMENT_BYTES: dict[str, bytes] = {}


def register_attachment_bytes(fixture: Fixture) -> None:
    """Remember the fixture's decoded attachment bytes by filename (test-local)."""
    for attachment in fixture.envelope["message"]["attachments"]:
        _ATTACHMENT_BYTES[attachment["filename"]] = base64.b64decode(attachment["content_b64"])


def _docx_text(data: bytes) -> str:
    doc = DocxDocument(io.BytesIO(data))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


async def run_fixture(rig: IntakeRig, fixture: Fixture, *, model_alias: str) -> IntakeResult:
    """Land → worker (prompt + run) → composition (live gateway) → read the record."""
    register_attachment_bytes(fixture)
    thread_id, project_id = await land(rig, fixture)
    if project_id is not None:
        await materialise_attachment_text(rig, project_id)

    started = await process_intake_thread(
        rig.factory, get_settings(), thread_id, enqueue=_always_enqueued
    )
    run_id: uuid.UUID | None = None
    run_status: str | None = None
    error: str | None = None
    tools_called: list[str] = []
    if started.get("status") == "started":
        run_id = uuid.UUID(started["run_id"])
        async with rig.factory() as db:
            run = await db.get(AgentRun, run_id)
            assert run is not None
            run.model_alias = model_alias
            await db.commit()
        registry = load_registry(SKILLS_DIR) if SKILLS_DIR.exists() else None
        await compose_and_execute_run(
            run_id=run_id,
            session_factory_provider=lambda: rig.factory,
            checkpointer_provider=InMemorySaver,
            skill_registry_provider=lambda: registry,
        )
        # In production this runs in the arq job's `finally` (agent_run_worker);
        # the rig drives composition directly, so it mirrors that one call here.
        await safe_fail_intake_thread(rig.factory, run_id)
        async with rig.factory() as db:
            run = await db.get(AgentRun, run_id)
            assert run is not None
            run_status, error = run.status, run.error
            steps = (
                (
                    await db.execute(
                        select(AgentRunStep)
                        .where(AgentRunStep.run_id == run_id)
                        .order_by(AgentRunStep.seq)
                    )
                )
                .scalars()
                .all()
            )
            tools_called = [s.name for s in steps if s.kind == "tool_call" and s.name]

    async with rig.factory() as db:
        thread = await db.get(IntakeThread, thread_id)
        assert thread is not None
        project = await db.get(Project, project_id) if project_id is not None else None
        return IntakeResult(
            fixture=fixture.name,
            thread_id=thread_id,
            project_id=project_id,
            run_id=run_id,
            run_status=run_status,
            outcome=thread.outcome,
            label=thread.label,
            thread_status=thread.status,
            project_archived=bool(project is not None and project.archived_at is not None),
            tools_called=tools_called,
            error=error,
        )


async def outbound_drafts(rig: IntakeRig, thread_id: uuid.UUID) -> list[IntakeMessage]:
    async with rig.factory() as db:
        return list(
            (
                await db.execute(
                    select(IntakeMessage).where(
                        IntakeMessage.thread_id == thread_id, IntakeMessage.direction == "out"
                    )
                )
            )
            .scalars()
            .all()
        )
