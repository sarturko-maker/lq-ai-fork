"""The Orion software development & licence agreement — a second COMPLEX
Commercial redline corpus instrument (C9).

A different dense instrument from Helios: the hard clauses here are an
**acceptance-testing procedure**, a **developed-IP / licence** structure, a
**support SLA with service credits**, and a layered **limitation of liability** —
each long, mostly sound, with only specific one-sided phrases. The surgical test
is the same: change the few words that need changing and leave the procedure,
the SLA mechanics and the existing carve-outs BARE; do not strike a whole clause
to fix one limb.

Same single-source contract (uploaded ``.docx`` text == searchable
``normalized_content``). The "leave-alone" traps:

- §2 Acceptance: a sound delivery → test-period → defect → re-delivery procedure
  that must stay BARE — the fix is the 5-day deemed-acceptance window, the
  pay-regardless line and adding a material-defect rejection right.
- §3 Licence & IP: a sound background-IP + licence-grant structure — the fix is
  the bespoke developments vesting only (assign or licence-back).
- §4 Support: a sound SLA (targets, response times, service credits) — the fix is
  the "sole and exclusive remedy" reach and a chronic-failure exit, not the SLA.
- §6 Limitation of liability: an existing carve-out limb that must stay BARE — add
  heads to it; do not rewrite the clause.
"""

from __future__ import annotations

import io

TITLE = "SOFTWARE DEVELOPMENT AND LICENCE AGREEMENT"
INTRO = (
    "This Software Development and Licence Agreement is entered into between Orion "
    'Software Labs Ltd (the "Developer") and the customer identified in the Order '
    '(the "Client"), under which the Developer will design, develop and license the '
    "software described in each Statement of Work and provide ongoing support."
)

CLAUSES: list[tuple[str, str]] = [
    (
        "1. Definitions",
        '"Software" means the software the Developer develops or licenses under a Statement of '
        'Work. "Deliverable" means any item delivered under a Statement of Work, including the '
        'Software and its documentation. "Client Materials" means materials the Client provides '
        'to the Developer. "Background IP" means intellectual property owned by a party before, '
        "or developed independently of, this Agreement.",
    ),
    (
        "2. Development and Acceptance",
        "The Developer shall deliver each Deliverable in accordance with the Statement of Work. "
        "On delivery the Client shall have a test period to evaluate the Deliverable against the "
        "agreed specification, during which it shall report any defects in writing with "
        "reasonable detail. The Developer shall correct reported defects and re-deliver, and the "
        "test period shall restart for the corrected Deliverable. A Deliverable shall be deemed "
        "accepted if the Client does not report any defect within five (5) days of delivery. The "
        "Client shall pay the milestone fee for each Deliverable on delivery whether or not the "
        "Deliverable has been accepted.",
    ),
    (
        "3. Licence and Intellectual Property",
        "Each party retains all right, title and interest in and to its own Background IP. The "
        "Developer grants the Client a non-exclusive, non-transferable, perpetual licence to use "
        "the Software for its internal business purposes. All right, title and interest in the "
        "Deliverables, including any bespoke developments commissioned and paid for by the "
        "Client, shall vest in and remain with the Developer. The Client grants the Developer a "
        "licence to use the Client Materials as necessary to perform the Statement of Work.",
    ),
    (
        "4. Support and Service Levels",
        "The Developer shall provide support for the Software during the term. The Developer shall "
        "use reasonable efforts to achieve 99.0% monthly availability and to respond to a "
        "priority-one incident within eight (8) business hours. If the Developer fails to meet a "
        "service level, the Client's sole and exclusive remedy shall be the service credits set "
        "out in the Statement of Work. The Developer shall maintain the support service "
        "substantially in accordance with its standard support policy.",
    ),
    (
        "5. Source Code Escrow",
        "The Developer shall deposit the source code for the Software with a reputable escrow "
        "agent and maintain it current. The escrow agent shall release the source code to the "
        "Client only if the Developer enters bankruptcy or liquidation. On release the Client may "
        "use the source code solely to maintain and support the Software for its internal use.",
    ),
    (
        "6. Indemnification",
        "The Developer shall defend the Client against any third-party claim that the Software as "
        "delivered infringes that third party's intellectual property rights, and shall indemnify "
        "the Client against damages finally awarded on such a claim. The Client shall indemnify, "
        "defend and hold harmless the Developer against any and all claims, losses and expenses "
        "arising from or in connection with the Client's use of the Software. The indemnified "
        "party shall give prompt written notice of the claim, allow the indemnifying party to "
        "control the defence, and provide reasonable cooperation at the indemnifying party's "
        "expense.",
    ),
    (
        "7. Limitation of Liability",
        "Except for the excluded matters set out below, neither party shall be liable for any "
        "indirect, incidental, special or consequential damages, or for any loss of profits, "
        "revenue or data, arising out of or in connection with this Agreement. Except for the "
        "excluded matters set out below, the Developer's total aggregate liability arising out of "
        "or in connection with this Agreement shall not exceed the fees paid by the Client in the "
        "three (3) months preceding the event giving rise to the claim. The excluded matters, to "
        "which no limitation in this section applies, are liability for death or personal injury "
        "caused by negligence and liability for fraud.",
    ),
    (
        "8. Term and Termination",
        "This Agreement continues until terminated. Either party may terminate for the other's "
        "material breach not cured within thirty (30) days of written notice. The Developer may "
        "terminate for convenience on sixty (60) days notice; the Client may not terminate for "
        "convenience.",
    ),
    (
        "9. Governing Law",
        "This Agreement is governed by the laws of England and Wales, and the parties submit to "
        "the exclusive jurisdiction of the courts of England and Wales.",
    ),
]

ORION_FILENAME = "Orion-Software-Development-Licence-Agreement.docx"


def _paragraphs() -> list[str]:
    out = [TITLE, INTRO]
    for heading, body in CLAUSES:
        out.append(heading)
        out.append(body)
    return out


def build_orion_docx() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(TITLE, level=0)
    doc.add_paragraph(INTRO)
    for heading, body in CLAUSES:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def orion_normalized_text() -> str:
    return "\n".join(_paragraphs())
