"""C4 live redline — does the Commercial agent produce a surgical tracked-changes
.docx on a real vendor-favoured MSA? (ADR-F031, provider-marked, CI-skipped.)

Run against the live dev stack (DeepSeek):

    DATABASE_URL=... LQ_AI_GATEWAY_URL=... LQ_AI_GATEWAY_KEY=... \\
    LQ_AI_SCENARIO_MODEL=deepseek UX_B1_EVIDENCE_DIR=<repo>/docs/fork/evidence/c4/live \\
    pytest -m provider tests/agents/scenarios/test_commercial_redline_scenario.py -s

It seeds a Commercial matter with the **real** SecureScan MSA `.docx` in object
storage (so `apply_redline` can fetch + redline the actual bytes), drives the
production agent loop, then captures what the model produced: the redlined
`.docx`, a readable reconstruction (`[-del-][+ins+]`), the accept-to-clean text,
and an adversarial **redline-quality judge** verdict (the §5.1 rubric, read
against the produced document). Per ADR-F015 the model's redline QUALITY is a
recorded finding, not a pass/fail gate — the hard assertions only confirm the
SYSTEM worked (the tool ran and emitted a tracked-changes work product).
"""

from __future__ import annotations

import json
import os
import uuid
from pathlib import Path

import pytest
import pytest_asyncio
from langchain_core.messages import HumanMessage, SystemMessage
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncEngine, AsyncSession, async_sessionmaker

from app import storage
from app.agents.factory import build_gateway_chat_model, build_gateway_http_client
from app.agents.redline_render import reconstruct_redline_text
from app.agents.redline_service import RedlineService
from app.models.document import Document, DocumentChunk
from app.models.file import File
from app.models.practice_area import PracticeArea
from app.models.project import Project
from app.models.user import User
from app.pipeline.readers._base import OOXML_DOCX_MIME
from app.security import hash_password
from tests.agents.scenarios.harness import run_scenario
from tests.agents.scenarios.scenarios import Scenario
from tests.agents.scenarios.securescan_msa import (
    MSA_FILENAME,
    build_msa_docx,
    msa_normalized_text,
)

pytestmark = [
    pytest.mark.provider,
    pytest.mark.skipif(
        "LQ_AI_GATEWAY_KEY" not in os.environ,
        reason="needs a live gateway (LQ_AI_GATEWAY_KEY unset)",
    ),
]

_EVIDENCE_DIR = (
    Path(os.environ["UX_B1_EVIDENCE_DIR"])
    if os.environ.get("UX_B1_EVIDENCE_DIR")
    else Path(__file__).resolve().parents[4] / "docs" / "fork" / "evidence" / "c4" / "live"
)
_MODEL = os.environ.get("LQ_AI_SCENARIO_MODEL", "deepseek")
_TERMINAL = {"completed", "failed", "cap_exceeded", "cancelled"}

REDLINE_SCENARIO = Scenario(
    id="commercial_redline_msa",
    title="Surgical redline of a vendor-favoured SaaS MSA",
    note="Does the agent read the MSA, then apply_redline with surgical, balanced edits?",
    prompt=(
        "You are our in-house commercial counsel and we are the CUSTOMER. Review the "
        f'attached vendor SaaS agreement "{MSA_FILENAME}" and redline it to protect us — '
        "apply your changes as tracked changes using apply_redline. Make surgical, "
        "balanced edits (carve-outs, deemed-direct, super-caps where appropriate); do not "
        "rip-and-replace whole clauses. Work through EVERY one-sided clause: the liability "
        "cap, the one-sided indemnity, the IP assignment, the Customer Data licence, the "
        "warranty disclaimer, the unilateral fee increase and the auto-renewal/no-termination "
        "term. Read the whole agreement first, then apply ALL your edits in a SINGLE "
        "apply_redline call (it takes a batch of many edits and re-redlines the original on "
        "each call, so one call must cover everything)."
    ),
    expect_tools=("apply_redline",),
    step_bound=100,
)


@pytest_asyncio.fixture
async def commit_factory(test_engine: AsyncEngine) -> async_sessionmaker[AsyncSession]:
    return async_sessionmaker(bind=test_engine, expire_on_commit=False, class_=AsyncSession)


async def _seed_msa_matter(
    factory: async_sessionmaker[AsyncSession],
) -> tuple[uuid.UUID, uuid.UUID]:
    """User + Commercial matter + the REAL MSA .docx in MinIO (File+Document+chunks)."""
    docx_bytes = build_msa_docx()
    text = msa_normalized_text()
    storage_key = str(uuid.uuid4())
    await storage.upload_bytes(
        storage_path=storage_key, body=docx_bytes, content_type=OOXML_DOCX_MIME
    )

    async with factory() as db:
        area_id = (
            await db.execute(select(PracticeArea.id).where(PracticeArea.key == "commercial"))
        ).scalar_one()
        user = User(
            email=f"c4-redline-{uuid.uuid4().hex[:8]}@example.com",
            display_name="C4 Redline User",
            hashed_password=hash_password("correct-horse-battery-staple"),
            is_admin=False,
            mfa_enabled=False,
            must_change_password=False,
        )
        db.add(user)
        await db.flush()
        project = Project(
            owner_id=user.id,
            name="SecureScan — Master Services Agreement",
            slug=f"commercial-{uuid.uuid4().hex[:6]}",
            privileged=True,
            minimum_inference_tier=4,
            practice_area_id=area_id,
        )
        db.add(project)
        await db.flush()
        file = File(
            owner_id=user.id,
            project_id=project.id,
            filename=MSA_FILENAME,
            mime_type=OOXML_DOCX_MIME,
            size_bytes=len(docx_bytes),
            hash_sha256=uuid.uuid4().hex,
            storage_path=storage_key,
            ingestion_status="ready",
        )
        db.add(file)
        await db.flush()
        document = Document(
            file_id=file.id,
            parser="c4-scenario-fixture",
            page_count=1,
            character_count=len(text),
            normalized_content=text,
        )
        db.add(document)
        await db.flush()
        # one chunk per paragraph so search_documents has something to match
        offset = 0
        for idx, para in enumerate(text.split("\n")):
            db.add(
                DocumentChunk(
                    document_id=document.id,
                    chunk_index=idx,
                    content=para,
                    page_start=1,
                    page_end=1,
                    char_offset_start=offset,
                    char_offset_end=offset + len(para),
                )
            )
            offset += len(para) + 1
        await db.commit()
        return user.id, project.id


async def _capture_redline(
    factory: async_sessionmaker[AsyncSession], user_id: uuid.UUID
) -> tuple[bytes, str] | None:
    """Find the redlined output File the agent produced; download its bytes."""
    async with factory() as db:
        row = (
            await db.execute(
                select(File.storage_path, File.filename)
                .where(File.owner_id == user_id, File.filename.like("%(redlined)%"))
                .order_by(File.created_at.desc())
            )
        ).first()
    if row is None:
        return None
    chunks: list[bytes] = []
    async with storage.stream_download(storage_path=row.storage_path) as stream:
        async for chunk in stream:
            chunks.append(chunk)
    return b"".join(chunks), row.filename


async def _judge(original_text: str, redline_view: str, accepted_text: str) -> str:
    """An adversarial redline-quality critic (the §5.1 rubric). Finding, not a gate."""
    http = build_gateway_http_client()
    try:
        model = build_gateway_chat_model(
            model_alias=_MODEL,
            purpose="commercial_redline_judge",
            http_async_client=http,
            project_minimum_inference_tier=None,
            privileged=False,
        )
        sys = SystemMessage(
            content=(
                "You are a senior commercial lawyer reviewing a JUNIOR's redline of a "
                "vendor-favoured SaaS MSA, acting for the CUSTOMER. Judge the redline on: "
                "(1) is it SURGICAL (narrow edits, not whole-clause rewrites)? (2) does it "
                "BALANCE one-sided clauses via the right mechanism (carve-outs of "
                "confidentiality/data/IP from the cap, deemed-direct losses, a super-cap, "
                "mutual indemnity, IP/licence-back) rather than just bumping numbers? "
                "(3) is it coherent across the agreement? (4) over- or under-protective? "
                "Give a verdict of STRONG / ADEQUATE / WEAK, then 3-6 bullet points. Be terse."
            )
        )
        # Generous caps so the judge sees the WHOLE redline (a 10-clause MSA
        # reconstruction is ~5 KB; truncating it produced a false-negative verdict).
        human = HumanMessage(
            content=(
                "ORIGINAL:\n" + original_text[:8000] + "\n\n"
                "THE REDLINE ([-deleted-][+inserted+]):\n" + redline_view[:16000] + "\n\n"
                "ACCEPTED RESULT:\n" + accepted_text[:10000]
            )
        )
        resp = await model.ainvoke([sys, human])
        return resp.content if isinstance(resp.content, str) else str(resp.content)
    finally:
        await http.aclose()


async def test_commercial_redline_live(
    commit_factory: async_sessionmaker[AsyncSession],
) -> None:
    user_id, _project_id = await _seed_msa_matter(commit_factory)
    _EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)
    (_EVIDENCE_DIR / MSA_FILENAME).write_bytes(build_msa_docx())

    receipt = await run_scenario(
        REDLINE_SCENARIO,
        _seeded(commit_factory, user_id, _project_id),
        model_alias=_MODEL,
        max_steps=100,
    )

    report: dict[str, object] = {"model": _MODEL, **receipt.to_dict()}
    captured = await _capture_redline(commit_factory, user_id)
    if captured is not None:
        redlined_bytes, redlined_name = captured
        (_EVIDENCE_DIR / redlined_name).write_bytes(redlined_bytes)
        redline_view = reconstruct_redline_text(redlined_bytes)
        accepted = _docx_text(RedlineService().accept_all(redlined_bytes))
        (_EVIDENCE_DIR / "redline-reconstruction.txt").write_text(redline_view, encoding="utf-8")
        (_EVIDENCE_DIR / "accepted-clean.txt").write_text(accepted, encoding="utf-8")
        report["redlined_file"] = redlined_name
        try:
            verdict = await _judge(msa_normalized_text(), redline_view, accepted)
            report["judge_verdict"] = verdict
            (_EVIDENCE_DIR / "judge-verdict.md").write_text(verdict, encoding="utf-8")
        except Exception as exc:
            report["judge_error"] = f"{type(exc).__name__}: {exc}"
    else:
        report["redlined_file"] = None

    (_EVIDENCE_DIR / "redline-report.json").write_text(
        json.dumps(report, indent=2), encoding="utf-8"
    )

    # Rig assertions only (ADR-F015): the loop turned the model and settled.
    assert receipt.status in _TERMINAL, receipt.status
    assert receipt.model_turns > 0, "the model never turned"
    # The system finding the maintainer reviews lives in the evidence dir.


# --- small local helpers (kept out of the harness) -------------------------- #


def _seeded(factory: async_sessionmaker[AsyncSession], user_id: uuid.UUID, project_id: uuid.UUID):
    from tests.agents.scenarios.harness import SeededMatter

    async def _noop_cleanup() -> None:
        return None

    return SeededMatter(
        factory=factory,
        user_id=user_id,
        project_id=project_id,
        practice_area_id=uuid.uuid4(),  # unused by run_scenario
        cleanup=_noop_cleanup,
    )


def _docx_text(data: bytes) -> str:
    import io

    from docx import Document as DocxDocument

    return "\n".join(p.text for p in DocxDocument(io.BytesIO(data)).paragraphs)
