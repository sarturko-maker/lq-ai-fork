"""The Northwind data processing agreement — a fourth Commercial redline corpus
instrument (C9), drafted for the Processor.

A controller-to-processor DPA skewed against the Controller: processing on the
Processor's own policies rather than documented instructions, unrestricted
sub-processing, a slow breach-notification window, no audit right, transfers at
the Processor's discretion, retention after termination, and — the head a
data-protection lawyer kills first — the Processor's right to use the personal
data for its own analytics and model training. A distinct instrument (a
schedule-style agreement, GDPR vocabulary) from the MSA/licence/NDA.

Same single-source contract (uploaded ``.docx`` text == searchable
``normalized_content``) with recognisable boilerplate (the Art. 32 measures
phrase) to keep bare and anchorable one-sided heads a Controller-side lawyer
rebalances.
"""

from __future__ import annotations

import io

TITLE = "DATA PROCESSING AGREEMENT"
INTRO = (
    "This Data Processing Agreement is entered into between the customer identified "
    'in the Order (the "Controller") and Northwind Cloud Services Ltd (the '
    '"Processor"), and governs the Processor\'s processing of personal data on behalf '
    "of the Controller in connection with the Services."
)

CLAUSES: list[tuple[str, str]] = [
    (
        "1. Definitions",
        '"Personal Data", "processing", "controller", "processor" and "data subject" have the '
        'meanings given in applicable data protection law. "Sub-processor" means any third party '
        "engaged by the Processor to process Personal Data.",
    ),
    (
        "2. Processing Instructions",
        "The Processor shall process the Personal Data in accordance with its own standard "
        "operating policies and as the Processor reasonably considers necessary to provide the "
        "Services, and may process the Personal Data for its own related business purposes.",
    ),
    (
        "3. Use of Data",
        "The Processor may use the Personal Data, including in de-identified or aggregated form, "
        "to develop, train and improve its products, analytics and machine-learning models, and "
        "to compile statistics, for its own account and without further obligation to the "
        "Controller.",
    ),
    (
        "4. Sub-processing",
        "The Processor may engage and replace Sub-processors at any time in its sole discretion "
        "and without prior notice to or authorisation from the Controller, and the Controller "
        "hereby consents to any such Sub-processor.",
    ),
    (
        "5. Security",
        "The Processor shall implement such technical and organisational measures as the Processor "
        "in its sole discretion considers appropriate, and shall not be required to implement any "
        "specific measure requested by the Controller.",
    ),
    (
        "6. Personal Data Breach",
        "The Processor shall notify the Controller of a personal data breach affecting the "
        "Controller's Personal Data within thirty (30) days of the Processor confirming the "
        "breach, and shall have no obligation to assist the Controller with its own notification "
        "duties.",
    ),
    (
        "7. Sub-processor Liability",
        "The Processor shall not be liable for any act or omission of any Sub-processor, and the "
        "Controller's sole recourse in respect of a Sub-processor shall be against that "
        "Sub-processor directly.",
    ),
    (
        "8. Audit",
        "The Controller shall have no right to audit or inspect the Processor's facilities, "
        "systems or records. The Processor may instead provide a summary certificate of "
        "compliance once per year, which the Controller agrees to accept as sufficient.",
    ),
    (
        "9. International Transfers",
        "The Processor may transfer the Personal Data to, and process it in, any country in which "
        "the Processor or its Sub-processors operate, at the Processor's discretion and without "
        "implementing any specific transfer safeguard.",
    ),
    (
        "10. Return and Deletion",
        "On termination the Processor shall delete the Personal Data where practicable, but may "
        "retain the Personal Data for as long as the Processor considers necessary for its own "
        "business, legal or analytics purposes.",
    ),
    (
        "11. Liability",
        "The Processor's aggregate liability arising out of or in connection with this Agreement "
        "and any processing of Personal Data shall not exceed the fees paid by the Controller in "
        "the one (1) month preceding the claim.",
    ),
    (
        "12. Governing Law",
        "This Agreement is governed by the laws of England and Wales, and the parties submit to "
        "the exclusive jurisdiction of the courts of England and Wales.",
    ),
]

DPA_FILENAME = "Northwind-DPA.docx"


def _paragraphs() -> list[str]:
    out = [TITLE, INTRO]
    for heading, body in CLAUSES:
        out.append(heading)
        out.append(body)
    return out


def build_dpa_docx() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(TITLE, level=0)
    doc.add_paragraph(INTRO)
    for heading, body in CLAUSES:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def dpa_normalized_text() -> str:
    return "\n".join(_paragraphs())
