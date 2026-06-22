"""The Aegis "mutual" NDA — a third Commercial redline corpus instrument (C9).

Labelled *mutual* but drafted one-directionally for the Discloser: an overbroad
definition with no standard exclusions, a perpetual confidentiality term, return
on demand with no backup/legal-retention carve-out, one-sided injunctive relief,
and a one-way indemnity. A different instrument from the MSA/licence so the
Claude-judged corpus spans contract types, with anchorable one-sided heads a
Receiving-Party-side lawyer rebalances and recognisable boilerplate to keep bare.

Same single-source contract (the uploaded ``.docx`` text == the searchable
``normalized_content``) so the agent's anchors match what Adeu finds in the runs.
"""

from __future__ import annotations

import io

TITLE = "MUTUAL NON-DISCLOSURE AGREEMENT"
INTRO = (
    "This Mutual Non-Disclosure Agreement is entered into between Aegis Analytics "
    'Inc. (the "Discloser") and the counterparty identified in the signature block '
    '(the "Recipient"), to protect Confidential Information exchanged in connection '
    "with a potential business relationship."
)

CLAUSES: list[tuple[str, str]] = [
    (
        "1. Definitions",
        '"Confidential Information" means all information disclosed by the Discloser to '
        "the Recipient, in any form and whether or not marked confidential, including all "
        "business, technical, financial and commercial information and any notes or analyses "
        "derived from it. All Confidential Information remains the property of the Discloser.",
    ),
    (
        "2. Confidentiality Obligations",
        "The Recipient shall hold in strict confidence all Confidential Information and shall "
        "not disclose it to any third party or use it for any purpose other than evaluating the "
        "potential business relationship, without the prior written consent of the Discloser.",
    ),
    (
        "3. Exclusions",
        "Information is not Confidential Information only to the extent the Recipient can show by "
        "written records that it was already in the public domain at the time of disclosure "
        "through no act of the Recipient.",
    ),
    (
        "4. Term",
        "The Recipient's obligations of confidentiality under this Agreement shall survive in "
        "perpetuity and shall continue to bind the Recipient indefinitely after the return or "
        "destruction of the Confidential Information.",
    ),
    (
        "5. Return of Information",
        "Upon the Discloser's written demand, the Recipient shall immediately return or destroy "
        "all Confidential Information and all copies, and shall certify such destruction in "
        "writing. The Recipient may not retain any copy for any reason.",
    ),
    (
        "6. No Licence",
        "Nothing in this Agreement grants the Recipient any licence or right in the Confidential "
        "Information, and the Discloser may pursue any opportunity with any third party without "
        "obligation to the Recipient.",
    ),
    (
        "7. Injunctive Relief",
        "The Recipient acknowledges that any breach would cause irreparable harm and agrees that "
        "the Discloser shall be entitled to injunctive relief without the requirement to post a "
        "bond or prove actual damages, in addition to any other remedy.",
    ),
    (
        "8. No Warranty",
        'All Confidential Information is provided "AS IS". The Discloser makes no representation '
        "or warranty as to its accuracy or completeness and shall have no liability to the "
        "Recipient arising from the Recipient's use of or reliance on it.",
    ),
    (
        "9. Indemnity",
        "The Recipient shall indemnify, defend and hold harmless the Discloser and its affiliates "
        "against any and all claims, losses, damages, liabilities and expenses arising from or in "
        "connection with the Recipient's handling of the Confidential Information.",
    ),
    (
        "10. Governing Law",
        "This Agreement is governed by the laws of the State of California, and the parties submit "
        "to the exclusive jurisdiction of the courts of California.",
    ),
]

NDA_FILENAME = "Aegis-Mutual-NDA.docx"


def _paragraphs() -> list[str]:
    out = [TITLE, INTRO]
    for heading, body in CLAUSES:
        out.append(heading)
        out.append(body)
    return out


def build_nda_docx() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(TITLE, level=0)
    doc.add_paragraph(INTRO)
    for heading, body in CLAUSES:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def nda_normalized_text() -> str:
    return "\n".join(_paragraphs())
