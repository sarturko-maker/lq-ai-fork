"""The Helios master SaaS & services agreement — a COMPLEX Commercial redline
corpus instrument (C9).

Unlike the short-clause MSA/licence/NDA, this instrument is built for the *hard*
surgical test the maintainer asked for: the operative clauses (limitation of
liability, indemnification, IP, warranties) are **long and multi-limb**, and most
of each clause is neutral, well-drafted language that a competent lawyer LEAVES
ALONE. Only specific phrases are one-sided. The correct redline is therefore a
handful of NARROW edits *inside* big blocks — striking a whole clause to fix one
limb would destroy good language (existing carve-outs, the indemnification
procedure, the exclusion list). It is exactly where a weaker model is tempted to
rip-and-replace.

Same single-source contract (uploaded ``.docx`` text == searchable
``normalized_content``). One-sided heads, and the surgical "leave-alone" traps:

- §6 Limitation of liability: a sound mutual exclusion-of-indirect list and an
  existing carve-out limb (death/PI, fraud) that must stay BARE — the fix is to
  ADD heads (confidentiality, data, IP) to the carve-out and extend the 6-month
  time-bar, not rewrite the clause.
- §5 Indemnification: a sound notice/defence/cooperation PROCEDURE that must stay
  BARE — the fix is to narrow the customer limb's trigger and add a settlement-
  consent proviso, not restate the procedure.
- §4 IP: background-IP and feedback limbs that are fine — the fix is the
  foreground/deliverables vesting only.
- §3 Warranties: a sound service warranty + remedy that is fine — the fix is the
  blanket disclaimer's reach.
"""

from __future__ import annotations

import io

TITLE = "MASTER SAAS AND SERVICES AGREEMENT"
INTRO = (
    "This Master SaaS and Services Agreement is entered into between Helios "
    'Platforms Inc. (the "Provider") and the customer identified in the Order Form '
    "(the \"Customer\"), and governs the Customer's access to the Provider's hosted "
    "platform and the professional services described in any Order Form or Statement "
    "of Work."
)

CLAUSES: list[tuple[str, str]] = [
    (
        "1. Definitions",
        '"Order Form" means an ordering document executed by the parties. "Services" means '
        'the Provider\'s hosted platform and any professional services ordered. "Customer Data" '
        "means data submitted by or on behalf of the Customer to the Services. "
        '"Confidential Information" means non-public information disclosed by a party that is '
        "marked confidential or that a reasonable person would understand to be confidential. "
        '"Deliverables" means items the Provider is engaged to create for the Customer under a '
        "Statement of Work.",
    ),
    (
        "2. Fees and Payment",
        "The Customer shall pay the fees set out in each Order Form. Undisputed invoices are "
        "payable within thirty (30) days of the invoice date, and the Customer may withhold any "
        "amount it disputes in good faith pending resolution. The Provider may increase the fees "
        "for any renewal term, and may also increase the fees during a term at any time in its "
        "sole discretion on thirty (30) days notice. Fees are exclusive of taxes, which the "
        "Customer is responsible for other than taxes on the Provider's income.",
    ),
    (
        "3. Warranties",
        "The Provider warrants that the Services will perform materially in accordance with the "
        "applicable documentation during the term, and that it will provide the professional "
        "services in a professional and workmanlike manner. If the Services do not conform, the "
        "Provider will use reasonable efforts to correct the non-conformity, and if it cannot do "
        "so within thirty (30) days the Customer may terminate the affected Order Form and "
        "receive a refund of prepaid, unused fees. EXCEPT FOR THE EXPRESS WARRANTIES IN THIS "
        "SECTION, THE SERVICES AND DELIVERABLES ARE PROVIDED AS IS AND THE PROVIDER DISCLAIMS "
        "ALL OTHER WARRANTIES, INCLUDING THE IMPLIED WARRANTIES OF MERCHANTABILITY, FITNESS FOR "
        "A PARTICULAR PURPOSE, NON-INFRINGEMENT, AND ANY WARRANTY THAT THE SERVICES WILL BE "
        "UNINTERRUPTED OR ERROR-FREE.",
    ),
    (
        "4. Intellectual Property",
        "As between the parties, each party retains all right, title and interest in and to its "
        "own background intellectual property, being intellectual property owned or developed "
        "independently of this Agreement. The Provider retains all right, title and interest in "
        "the Services and its platform. All right, title and interest in the Deliverables, and in "
        "any configurations and customisations developed under a Statement of Work, shall vest "
        "exclusively in the Provider on creation, and the Customer receives only a non-exclusive "
        "licence to use them for its internal business purposes. Each party may use any general "
        "skills, know-how and feedback arising from the engagement.",
    ),
    (
        "5. Indemnification",
        "The Provider shall defend the Customer against any third-party claim that the Services "
        "as provided infringe a patent, copyright or trade secret of that third party, and shall "
        "indemnify the Customer against damages finally awarded. The Customer shall indemnify, "
        "defend and hold harmless the Provider against any and all claims, losses and expenses "
        "arising from or in connection with the Customer's use of the Services. The party seeking "
        "indemnity shall give the indemnifying party prompt written notice of the claim, shall "
        "permit the indemnifying party to control the defence and settlement of the claim, and "
        "shall provide reasonable cooperation at the indemnifying party's expense; provided that "
        "the indemnifying party may settle any claim in its sole discretion.",
    ),
    (
        "6. Limitation of Liability",
        "Except for the excluded matters described below, neither party shall be liable for any "
        "indirect, incidental, special, consequential or punitive damages, or for any loss of "
        "profits, revenue, goodwill or anticipated savings, arising out of or in connection with "
        "this Agreement, even if advised of the possibility of such damages. Except for the "
        "excluded matters described below, each party's total aggregate liability arising out of "
        "or in connection with this Agreement shall not exceed the total fees paid or payable by "
        "the Customer under the applicable Order Form in the six (6) months preceding the event "
        "giving rise to the claim. The excluded matters, to which no limitation or exclusion in "
        "this section applies, are liability for death or personal injury caused by negligence "
        "and liability for fraud or fraudulent misrepresentation. No claim may be brought under "
        "this Agreement more than six (6) months after the date on which the cause of action "
        "first arose.",
    ),
    (
        "7. Term and Termination",
        "This Agreement continues for the term stated in the Order Form and renews automatically "
        "for successive terms of equal length unless either party gives sixty (60) days written "
        "notice of non-renewal. Either party may terminate for the other's material breach not "
        "cured within thirty (30) days of written notice. On expiry or termination the Customer "
        "may export Customer Data for thirty (30) days, after which the Provider may delete it.",
    ),
    (
        "8. Governing Law",
        "This Agreement is governed by the laws of the State of New York, and the parties submit "
        "to the exclusive jurisdiction of the courts located in New York County, New York.",
    ),
]

HELIOS_FILENAME = "Helios-Master-SaaS-Services-Agreement.docx"


def _paragraphs() -> list[str]:
    out = [TITLE, INTRO]
    for heading, body in CLAUSES:
        out.append(heading)
        out.append(body)
    return out


def build_helios_docx() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(TITLE, level=0)
    doc.add_paragraph(INTRO)
    for heading, body in CLAUSES:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def helios_normalized_text() -> str:
    return "\n".join(_paragraphs())
