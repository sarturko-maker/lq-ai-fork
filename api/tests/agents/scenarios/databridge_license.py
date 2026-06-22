"""The DataBridge vendor-favoured software licence + support agreement — a second
Commercial redline-eval input (C8), a different instrument from the SaaS MSA so the
surgical-craft eval has corpus breadth.

Same single-source contract (the uploaded ``.docx`` text == the searchable
``normalized_content``) and the same one-sided heads an in-house customer-side
lawyer rebalances (§5.1): licence scope, fees, IP, indemnity, warranty, liability,
termination — worded differently from SecureScan so it is a genuinely distinct test.
"""

from __future__ import annotations

import io

TITLE = "SOFTWARE LICENCE AND SUPPORT AGREEMENT"
INTRO = (
    "This Software Licence and Support Agreement is entered into between DataBridge "
    'Systems LLC (the "Licensor") and the customer identified in the Order (the '
    '"Licensee").'
)

CLAUSES: list[tuple[str, str]] = [
    (
        "1. Definitions",
        '"Software" means the DataBridge platform licensed under the Order. "Licensor" means '
        'DataBridge Systems LLC. "Licensee" means the customer identified in the Order. '
        '"Licensee Content" means data the Licensee loads into the Software.',
    ),
    (
        "2. Licence",
        "The Licensor grants the Licensee a non-exclusive, non-transferable licence to use the "
        "Software during the term. The Licensor may suspend access at any time in its sole "
        "discretion and without notice.",
    ),
    (
        "3. Fees",
        "The Licensee shall pay all fees within fifteen (15) days of the invoice date. The "
        "Licensor may increase the fees at any time on written notice. All fees are "
        "non-refundable under all circumstances.",
    ),
    (
        "4. Intellectual Property",
        "All intellectual property rights in any configurations, integrations and feedback "
        "created in connection with the Software shall vest exclusively in the Licensor upon "
        "creation.",
    ),
    (
        "5. Licensee Content",
        "The Licensee grants the Licensor a perpetual, irrevocable, worldwide, royalty-free "
        "licence to use, reproduce and create derivative works from Licensee Content for any "
        "purpose, including to train the Licensor's models and improve its products.",
    ),
    (
        "6. Warranties",
        'THE SOFTWARE IS PROVIDED "AS IS" AND THE LICENSOR DISCLAIMS ALL WARRANTIES, WHETHER '
        "EXPRESS, IMPLIED OR STATUTORY, INCLUDING MERCHANTABILITY, FITNESS FOR A PARTICULAR "
        "PURPOSE AND NON-INFRINGEMENT.",
    ),
    (
        "7. Indemnity",
        "The Licensee shall indemnify, defend and hold harmless the Licensor and its officers "
        "against any and all claims, losses, damages and expenses arising from or in connection "
        "with the Licensee's use of the Software or the Licensee Content.",
    ),
    (
        "8. Limitation of Liability",
        "The Licensor's total liability arising out of or in connection with this Agreement "
        "shall not exceed the fees paid by the Licensee in the three (3) months preceding the "
        "claim. In no event shall the Licensor be liable for any indirect, incidental or "
        "consequential damages, or any loss of profits, revenue or data.",
    ),
    (
        "9. Term and Termination",
        "This Agreement renews automatically for successive one (1) year terms. The Licensee "
        "may not terminate for convenience. The Licensor may terminate at any time on thirty "
        "(30) days notice.",
    ),
    (
        "10. Governing Law",
        "This Agreement is governed by the laws of the State of New York, and the parties submit "
        "to the exclusive jurisdiction of the courts of New York.",
    ),
]

LICENSE_FILENAME = "DataBridge-Licence.docx"


def _paragraphs() -> list[str]:
    out = [TITLE, INTRO]
    for heading, body in CLAUSES:
        out.append(heading)
        out.append(body)
    return out


def build_license_docx() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(TITLE, level=0)
    doc.add_paragraph(INTRO)
    for heading, body in CLAUSES:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def license_normalized_text() -> str:
    return "\n".join(_paragraphs())
