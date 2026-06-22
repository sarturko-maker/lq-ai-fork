"""The SecureScan vendor-favoured SaaS MSA — the C4 live-redline scenario input.

Single source so the uploaded ``.docx`` and the searchable ``normalized_content``
share identical text (the agent quotes anchors from what it reads; Adeu must find
them in the ``.docx`` runs). Deliberately one-sided in the heads of terms an
in-house customer-side lawyer rebalances (§5.1): liability cap, indemnity,
IP/data, warranty, termination, fees.
"""

from __future__ import annotations

import io

TITLE = "MASTER SERVICES AGREEMENT"
INTRO = (
    "This Master Services Agreement is entered into between SecureScan, Inc. (the "
    '"Vendor") and the Customer identified in the Order Form.'
)

CLAUSES: list[tuple[str, str]] = [
    (
        "1. Definitions",
        '"Agreement" means this Master Services Agreement. "Vendor" means SecureScan, Inc. '
        '"Customer" means the entity identified in the Order Form. "Services" means the '
        'Vendor\'s hosted SecureScan platform. "Customer Data" means data submitted by the '
        "Customer to the Services.",
    ),
    (
        "2. Services",
        "The Vendor shall use commercially reasonable efforts to make the Services available, "
        "but does not warrant uninterrupted or error-free operation. The Vendor may modify or "
        "discontinue any feature of the Services at any time in its sole discretion.",
    ),
    (
        "3. Fees",
        "The Customer shall pay all fees set out in the Order Form within thirty (30) days of "
        "the invoice date. All fees are non-refundable. The Vendor may increase fees at any "
        "time on thirty (30) days notice.",
    ),
    (
        "4. Term and Renewal",
        "This Agreement commences on the Effective Date and continues for an initial term of "
        "three (3) years. It shall automatically renew for successive three (3) year terms "
        "unless the Customer gives at least ninety (90) days written notice before the end of "
        "the then-current term. The Customer may not terminate for convenience.",
    ),
    (
        "5. Intellectual Property",
        "All intellectual property rights in any deliverables, configurations, feedback and "
        "work product created under this Agreement shall vest exclusively in the Vendor upon "
        "creation. The Customer assigns to the Vendor all right, title and interest in any "
        "feedback it provides.",
    ),
    (
        "6. Customer Data",
        "The Customer grants the Vendor a perpetual, irrevocable, worldwide, royalty-free "
        "licence to use, reproduce and create derivative works from Customer Data for any "
        "purpose, including to improve the Vendor's products and services.",
    ),
    (
        "7. Warranties",
        'THE SERVICES ARE PROVIDED "AS IS" AND THE VENDOR DISCLAIMS ALL WARRANTIES, WHETHER '
        "EXPRESS, IMPLIED OR STATUTORY, INCLUDING ANY WARRANTY OF MERCHANTABILITY, FITNESS FOR "
        "A PARTICULAR PURPOSE AND NON-INFRINGEMENT.",
    ),
    (
        "8. Indemnity",
        "The Customer shall indemnify, defend and hold harmless the Vendor and its affiliates "
        "against any and all claims, losses, damages, liabilities and expenses arising from or "
        "in connection with the Customer's use of the Services or the Customer Data.",
    ),
    (
        "9. Limitation of Liability",
        "The Vendor's aggregate liability arising out of or in connection with this Agreement "
        "shall not exceed the total fees paid by the Customer in the one (1) month preceding "
        "the claim. In no event shall the Vendor be liable for any indirect, incidental, "
        "special or consequential damages, or for any loss of profits, revenue or data, even "
        "if advised of the possibility of such damages.",
    ),
    (
        "10. Governing Law",
        "This Agreement is governed by the laws of the State of Delaware, and the parties "
        "submit to the exclusive jurisdiction of the courts of Delaware.",
    ),
]

MSA_FILENAME = "SecureScan-MSA.docx"


def _paragraphs() -> list[str]:
    out = [TITLE, INTRO]
    for heading, body in CLAUSES:
        out.append(heading)
        out.append(body)
    return out


def build_msa_docx() -> bytes:
    """The .docx as python-docx would author it (headings + body paragraphs)."""
    from docx import Document

    doc = Document()
    doc.add_heading(TITLE, level=0)
    doc.add_paragraph(INTRO)
    for heading, body in CLAUSES:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def msa_normalized_text() -> str:
    """Searchable/readable text matching the .docx paragraph order."""
    return "\n".join(_paragraphs())
