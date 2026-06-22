"""The Meridian professional-services agreement (statement of work) — a fifth
Commercial redline corpus instrument (C9), drafted for the Supplier.

A consulting / professional-services agreement skewed against the Customer:
time-and-materials with no cap and non-binding estimates, deemed acceptance,
deliverable IP vesting in the Supplier, no warranty that deliverables meet the
specification, free substitution of personnel, a one-way indemnity, a one-month
liability cap, and termination for convenience available only to the Supplier. A
distinct instrument (services, not a product licence) so the Claude-judged corpus
spans contract types.

Same single-source contract (uploaded ``.docx`` text == searchable
``normalized_content``) with recognisable boilerplate (the indemnity verb phrase
and the cap stem) to keep bare and anchorable one-sided heads a Customer-side
lawyer rebalances.
"""

from __future__ import annotations

import io

TITLE = "PROFESSIONAL SERVICES AGREEMENT"
INTRO = (
    "This Professional Services Agreement is entered into between Meridian Consulting "
    'Group LLC (the "Supplier") and the customer identified in the Statement of Work '
    '(the "Customer"), under which the Supplier will provide the consulting services '
    "and deliverables described in each Statement of Work."
)

CLAUSES: list[tuple[str, str]] = [
    (
        "1. Services",
        "The Supplier shall provide the services described in each Statement of Work on a "
        "time-and-materials basis. Any timelines, estimates or staffing levels are indicative "
        "only and do not bind the Supplier.",
    ),
    (
        "2. Acceptance",
        "Each deliverable shall be deemed accepted unless the Customer gives written notice of "
        "rejection within three (3) days of delivery. The Customer shall pay for all deliverables "
        "whether or not accepted, and the Supplier is not obliged to correct any deliverable at "
        "its own cost.",
    ),
    (
        "3. Fees and Expenses",
        "The Customer shall pay all fees on a time-and-materials basis without any cap, together "
        "with all expenses at cost, within fifteen (15) days of the invoice date. The Supplier "
        "may revise its rates at any time on written notice. All fees are non-refundable.",
    ),
    (
        "4. Intellectual Property",
        "All intellectual property rights in any deliverables, work product, methodologies and "
        "materials created under this Agreement shall vest exclusively in the Supplier. The "
        "Customer is granted a non-exclusive, non-transferable licence to use the deliverables "
        "for its internal business purposes only.",
    ),
    (
        "5. Warranties",
        'The services and deliverables are provided "AS IS". The Supplier does not warrant that '
        "any deliverable will conform to any specification, be fit for any particular purpose, or "
        "be free from defects, and disclaims all warranties to the fullest extent permitted by law.",
    ),
    (
        "6. Personnel",
        "The Supplier may substitute, reassign or remove any personnel at any time in its sole "
        "discretion, and gives no commitment that any named individual will perform the services.",
    ),
    (
        "7. Indemnity",
        "The Customer shall indemnify, defend and hold harmless the Supplier and its personnel "
        "against any and all claims, losses, damages, liabilities and expenses arising from or in "
        "connection with the services, the deliverables or the Customer's use of them.",
    ),
    (
        "8. Limitation of Liability",
        "The Supplier's aggregate liability arising out of or in connection with this Agreement "
        "shall not exceed the fees paid by the Customer in the one (1) month preceding the claim. "
        "In no event shall the Supplier be liable for any indirect, incidental or consequential "
        "damages, or for any loss of profits, revenue or data.",
    ),
    (
        "9. Term and Termination",
        "This Agreement continues until each Statement of Work is completed. The Supplier may "
        "terminate for convenience at any time on ten (10) days notice; the Customer may not "
        "terminate for convenience. On termination the Customer shall pay all fees for work in "
        "progress and for the full committed engagement.",
    ),
    (
        "10. Governing Law",
        "This Agreement is governed by the laws of the State of Illinois, and the parties submit "
        "to the exclusive jurisdiction of the courts of Illinois.",
    ),
]

SOW_FILENAME = "Meridian-Professional-Services-Agreement.docx"


def _paragraphs() -> list[str]:
    out = [TITLE, INTRO]
    for heading, body in CLAUSES:
        out.append(heading)
        out.append(body)
    return out


def build_sow_docx() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(TITLE, level=0)
    doc.add_paragraph(INTRO)
    for heading, body in CLAUSES:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def sow_normalized_text() -> str:
    return "\n".join(_paragraphs())
