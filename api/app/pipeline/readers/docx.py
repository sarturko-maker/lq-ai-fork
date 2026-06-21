"""DOCX reader — one unit per paragraph block, via python-docx (MIT).

A ``.docx`` has no reliable page boundaries, so the unit is the
paragraph block: ``page_number`` is a 1-based paragraph ordinal (a
deliberate reinterpretation of the ``page`` field for paginationless
formats — see ADR-F029). Text comes from ``document.paragraphs`` in body
order. **Known limitation (C1):** table-cell text is not extracted —
python-docx's ``paragraphs`` excludes tables and a body-order walk of
paragraphs+tables is deferred; flagged in ADR-F029.

The container is hardened by :func:`guard_ooxml` *before* python-docx
(lxml) opens it, so an entity-bearing document is rejected before any
XML parse.
"""

from __future__ import annotations

from io import BytesIO

from app.pipeline.readers._base import (
    OOXML_DOCX_MIME,
    ParsedDocument,
    ParserError,
    build_parsed_document,
    dist_version,
    guard_ooxml,
    ooxml_subtype,
)


class DocxReader:
    """Extracts paragraph-block text from a ``.docx`` document."""

    parser_label = "python-docx"
    mimes = frozenset({OOXML_DOCX_MIME})

    def sniff(self, data: bytes) -> bool:
        return ooxml_subtype(data) == "docx"

    def read(self, data: bytes) -> ParsedDocument:
        guard_ooxml(data)
        try:
            import docx
        except ImportError as exc:  # pragma: no cover - install-time error
            raise ParserError("python-docx is not installed; cannot read DOCX") from exc

        # The container passed sniff + guard_ooxml; a still-malformed document
        # must fail closed as ParserError (-> ingestion parse_failed), not
        # escape as a bare python-docx error (mirrors parsers._run_pymupdf).
        try:
            document = docx.Document(BytesIO(data))
            unit_texts = [paragraph.text for paragraph in document.paragraphs]
        except Exception as exc:
            raise ParserError(f"python-docx failed to read DOCX: {exc}") from exc

        return build_parsed_document(
            unit_texts,
            parser=self.parser_label,
            parser_version=f"python-docx={dist_version('python-docx')}",
        )
